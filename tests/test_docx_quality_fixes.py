"""2026-07 DOCX 排版质量修复的回归测试。

覆盖四个已确认 bug：
① 有序列表编号被嵌套无序子列表重置
② 行内/块级公式图片按位置配对（不再对调）+ 行内公式保持段内
③ 普通本地图片：bundle 资源嵌入 + 缺资源时 warning（不再静默）
④ 表格分隔行 :---: 对齐标记生效
"""
import base64
import io
import tempfile
from pathlib import Path

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image as PILImage

from app.generator import generate_docx_from_content, parse_markdown, BlockType, _parse_table_aligns
from app.main import (
    _embed_local_image_refs,
    _math_slots_in_document_order,
    _replace_katex_images_by_position,
)
from app.image_injector import preprocess_markdown


def _make_png_base64(width=40, height=20, color=(200, 30, 30)):
    img = PILImage.new("RGB", (width, height), color)
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return base64.b64encode(buf.getvalue()).decode("ascii")


def _render_docx(md: str, style: str = "standard") -> Document:
    with tempfile.TemporaryDirectory() as tmp:
        path = str(Path(tmp) / "out.docx")
        generate_docx_from_content(md, path, style=style)
        return Document(path)


def _paragraph_texts(doc: Document) -> list[str]:
    return [p.text for p in doc.paragraphs]


# ── ① 有序列表编号 ──────────────────────────────────────────────

class TestOrderedListNumbering:
    def test_nested_unordered_does_not_reset_parent_counter(self):
        md = "\n".join([
            "1. 第一项",
            "2. 第二项",
            "   - 子要点甲",
            "   - 子要点乙",
            "3. 第三项",
        ])
        texts = _paragraph_texts(_render_docx(md))
        joined = "\n".join(texts)
        assert any(t.startswith("3.") and "第三项" in t for t in texts), f"第三项编号应为 3.：{joined}"
        assert not any(t.startswith("1.") and "第三项" in t for t in texts)

    def test_sibling_lists_still_restart(self):
        md = "\n".join([
            "1. 甲",
            "",
            "中间正文段落。",
            "",
            "1. 乙",
        ])
        texts = _paragraph_texts(_render_docx(md))
        numbered = [t for t in texts if t.startswith("1.")]
        assert len(numbered) == 2, f"两个独立列表都应从 1. 开始：{texts}"


# ── ② 公式位置配对 ──────────────────────────────────────────────

class TestKatexPositionPairing:
    def test_slots_enumerated_in_document_order(self):
        md = "开头 $inline$ 中间\n\n$$block$$\n\n结尾"
        slots = _math_slots_in_document_order(md)
        assert [s["kind"] for s in slots] == ["inline", "block"]

    def test_inline_inside_block_not_double_counted(self):
        md = "$$a + b$$"
        slots = _math_slots_in_document_order(md)
        assert [s["kind"] for s in slots] == ["block"]

    def test_images_paired_by_position_not_by_kind_preference(self):
        md = "得分按 $S$ 加权：\n\n$$\\sum w_i = 1$$"
        images = [
            {"id": "mdv__chart__aaaaaaaa__", "type": "katex", "widthPx": 100, "widthCm": 2.8},
            {"id": "mdv__chart__bbbbbbbb__", "type": "katex", "widthPx": 500, "widthCm": 9.0},
        ]
        result = _replace_katex_images_by_position(md, images)
        # 第一张图（渲染顺序=文档顺序）应落在行内位置，第二张落在块级位置
        inline_pos = result.find("mdv__chart__aaaaaaaa__")
        block_pos = result.find("mdv__chart__bbbbbbbb__")
        assert 0 < inline_pos < block_pos, result
        assert images[0].get("inline") is True
        assert "inline" not in images[1]
        # 行内占位符不另起段落
        assert "得分按 ![](mdv__chart__aaaaaaaa__) 加权：" in result

    def test_inline_placeholder_not_blockified_in_preprocess(self):
        png = _make_png_base64()
        images = [
            {"id": "mdv__chart__aaaaaaaa__", "pngBase64": png, "widthCm": 1.0, "inline": True},
            {"id": "mdv__chart__bbbbbbbb__", "pngBase64": png, "widthCm": 9.0},
        ]
        md = "正文 ![](mdv__chart__aaaaaaaa__) 继续\n\n![](mdv__chart__bbbbbbbb__)"
        out_md, image_map = preprocess_markdown(md, images)
        assert "正文 ![](mdv__chart__aaaaaaaa__) 继续" in out_md
        assert image_map["mdv__chart__aaaaaaaa__"].inline is True
        assert image_map["mdv__chart__bbbbbbbb__"].inline is False

    def test_fenced_code_dollars_are_not_math_slots(self):
        """bash 围栏里的 ${VAR}...${VAR} 不是公式槽位（真实运维文档踩雷回归）。"""
        md = "```bash\ngzip > ${BACKUP_DIR}/mysql_full_${DATE}.sql.gz\n```\n\n正文公式 $S$ 收尾。"
        slots = _math_slots_in_document_order(md)
        assert [s["kind"] for s in slots] == ["inline"]
        assert md[slots[0]["start"]:slots[0]["end"]] == "$S$"

    def test_sync_renderer_does_not_eat_dollars_in_fences(self):
        from app.chart_renderers import render_charts_and_formulas_sync
        md = "```bash\ntar -czf ${BACKUP_DIR}/x_$(date +%s).tar.gz -C /tmp $(basename ${TEMP_DIR})\n```\n"
        result = render_charts_and_formulas_sync(md, [])
        assert result.markdown == md
        assert result.images == {}

    def test_sync_renderer_still_renders_prose_math(self):
        from app.chart_renderers import render_charts_and_formulas_sync
        md = "```bash\necho ${A} ${B}\n```\n\n真公式 $a+b$ 在此。"
        result = render_charts_and_formulas_sync(md, [])
        assert "${A} ${B}" in result.markdown
        assert "$a+b$" not in result.markdown
        assert len(result.images) == 1

    def test_inline_code_dollars_are_not_math_slots(self):
        """表格单元格里反引号包的正则（含成对 $）不是公式槽位（K8s 实档撕表回归）。"""
        md = "| App | `v1/(chat(?!/completions$).*\\|tools/websearch$)` |\n\n正文 $S$ 收尾。"
        slots = _math_slots_in_document_order(md)
        assert [s["kind"] for s in slots] == ["inline"]
        assert md[slots[0]["start"]:slots[0]["end"]] == "$S$"

    def test_sync_renderer_does_not_eat_dollars_in_inline_code(self):
        from app.chart_renderers import render_charts_and_formulas_sync
        md = "| App | 30063 | `v1/(chat(?!/completions$).*\\|app/curr\\|websearch$)` |\n"
        result = render_charts_and_formulas_sync(md, [])
        assert result.markdown == md
        assert result.images == {}

    def test_tilde_fence_dollars_are_protected(self):
        """~~~ 围栏与 ``` 围栏同权：内部 $ 不是公式槽位。"""
        md = "~~~bash\ngzip > ${A}/x_${B}.gz\n~~~\n\n正文 $S$。"
        slots = _math_slots_in_document_order(md)
        assert [s["kind"] for s in slots] == ["inline"]
        assert md[slots[0]["start"]:slots[0]["end"]] == "$S$"

    def test_four_backtick_fence_protected(self):
        md = "````text\necho ${A} ${B}\n````\n"
        assert _math_slots_in_document_order(md) == []

    def test_double_backtick_inline_code_protected(self):
        from app.chart_renderers import render_charts_and_formulas_sync
        md = "规则 ``price$1 to $9`` 说明。\n"
        result = render_charts_and_formulas_sync(md, [])
        assert result.markdown == md
        assert result.images == {}

    def test_leftover_math_kept_when_more_slots_than_images(self):
        md = "$a$ 和 $b$"
        images = [{"id": "mdv__chart__aaaaaaaa__", "type": "katex", "widthPx": 80, "widthCm": 2.8}]
        result = _replace_katex_images_by_position(md, images)
        assert "mdv__chart__aaaaaaaa__" in result
        assert "$b$" in result  # 没有图的槽位保留原文，交给后续渲染轮次


# ── ③ 本地图片嵌入与告警 ─────────────────────────────────────────

class TestLocalImageEmbedding:
    def test_missing_resource_warns_instead_of_silence(self):
        md = "# 标题\n\n![示意图](./missing.png)\n"
        out_md, images, warnings = _embed_local_image_refs(md, [], None)
        assert images == []
        assert out_md == md  # 保留原文
        assert any("missing.png" in w for w in warnings)

    def test_bundle_resource_embedded(self):
        png = _make_png_base64(width=96, height=48)
        resources = [{
            "path": "assets/pic.png", "kind": "binary", "base64": png,
            "content": None, "mediaType": "image/png", "size": 100,
        }]
        md = "![配图](./assets/pic.png)\n"
        out_md, images, warnings = _embed_local_image_refs(md, resources, "doc.md")
        assert len(images) == 1
        assert warnings == []
        assert images[0]["pngBase64"] == png
        assert images[0]["widthPx"] == 96
        assert f"![]({images[0]['id']})" in out_md
        assert "assets/pic.png" not in out_md

    def test_entry_dir_relative_resolution(self):
        png = _make_png_base64()
        resources = [{
            "path": "docs/img/a.png", "kind": "binary", "base64": png,
            "content": None, "mediaType": "image/png", "size": 100,
        }]
        md = "![a](img/a.png)"
        out_md, images, warnings = _embed_local_image_refs(md, resources, "docs/readme.md")
        assert len(images) == 1 and warnings == []

    def test_remote_and_data_uris_ignored(self):
        md = "![r](https://example.com/a.png)\n\n![d](data:image/png;base64,xxxx)\n"
        out_md, images, warnings = _embed_local_image_refs(md, [], None)
        assert images == [] and warnings == [] and out_md == md

    def test_image_ref_inside_code_fence_untouched(self):
        md = "```text\n![展示](./demo.png)\n```\n"
        out_md, images, warnings = _embed_local_image_refs(md, [], None)
        assert images == [] and warnings == [] and out_md == md

    def test_unsupported_extension_warns(self):
        md = "![v](./chart.svg)\n"
        _, images, warnings = _embed_local_image_refs(md, [], None)
        assert images == []
        assert any("chart.svg" in w for w in warnings)

    def test_traversal_path_rejected_with_warning(self):
        md = "![x](../../etc/secret.png)\n"
        _, images, warnings = _embed_local_image_refs(md, [], None)
        assert images == []
        assert any("secret.png" in w for w in warnings)


# ── ④ 表格列对齐 ────────────────────────────────────────────────

class TestEscapedPipeInTableCells:
    """K8s 实档回归：单元格里的 \\| 转义竖线不许炸列。"""

    def test_escaped_pipes_stay_in_cell(self):
        from app.generator import _split_table_cells
        row = r"| App | 30063 | `v1/(chat\|app/curr\|websearch$)` | 说明 |"
        cells = _split_table_cells(row)
        assert len(cells) == 4
        assert cells[2] == "`v1/(chat|app/curr|websearch$)`"

    def test_parse_markdown_table_with_escaped_pipes(self):
        md = "| 名称 | 规则 |\n|------|------|\n| App | `a\\|b\\|c` |\n"
        blocks = [b for b in parse_markdown(md) if b.type == BlockType.TABLE]
        assert len(blocks) == 1
        assert blocks[0].rows[1] == ["App", "`a|b|c`"]

    def test_docx_table_keeps_columns(self):
        md = "| 名称 | 规则 |\n|------|------|\n| App | `a\\|b\\|c` |\n"
        doc = _render_docx(md)
        table = doc.tables[0]
        assert len(table.columns) == 2
        assert "a|b|c" in table.cell(1, 1).text


class TestTableColumnAlignment:
    def test_parse_aligns(self):
        assert _parse_table_aligns("| :--- | :---: | ---: | --- |") == ["left", "center", "right", None]

    def test_parse_markdown_keeps_aligns(self):
        md = "| 名称 | 数量 |\n|:-----|-----:|\n| 甲 | 12 |\n"
        blocks = [b for b in parse_markdown(md) if b.type == BlockType.TABLE]
        assert len(blocks) == 1
        assert blocks[0].column_aligns == ["left", "right"]

    @pytest.mark.parametrize("style", ["preview", "standard", "official"])
    def test_right_align_applied_in_docx(self, style):
        md = "| 名称 | 数量 |\n|:-----|-----:|\n| 甲 | 12 |\n| 乙 | 345 |\n"
        doc = _render_docx(md, style=style)
        tables = doc.tables
        assert tables, "应生成表格"
        cell = tables[0].cell(1, 1)  # 数据行第二列
        assert cell.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT
        left_cell = tables[0].cell(1, 0)
        assert left_cell.paragraphs[0].alignment in (WD_ALIGN_PARAGRAPH.LEFT, None)
