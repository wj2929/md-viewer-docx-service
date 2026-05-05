import base64
import io
import re
import pytest
from PIL import Image
from docx import Document

from app.image_injector import (
    ImageData, PLACEHOLDER_PATTERN, IMAGE_MAX_B64_LEN, IMAGE_MAX_PIXELS,
    preprocess_markdown, inject_images, ImageLayout,
)


class TestPlaceholderPattern:
    def test_valid_placeholder(self):
        assert PLACEHOLDER_PATTERN.match("![](mdv__chart__deadbeef__)")

    def test_valid_placeholder_extracts_id(self):
        m = PLACEHOLDER_PATTERN.match("![](mdv__chart__a1b2c3d4__)")
        assert m.group(1) == "mdv__chart__a1b2c3d4__"

    def test_rejects_wrong_prefix(self):
        assert PLACEHOLDER_PATTERN.match("![](chart__deadbeef__)") is None

    def test_rejects_short_hex(self):
        assert PLACEHOLDER_PATTERN.match("![](mdv__chart__dead__)") is None

    def test_rejects_non_hex(self):
        assert PLACEHOLDER_PATTERN.match("![](mdv__chart__ghijklmn__)") is None

    def test_rejects_extra_text(self):
        assert PLACEHOLDER_PATTERN.match("text ![](mdv__chart__deadbeef__)") is None


class TestImageData:
    def test_valid_png(self, minimal_png_base64):
        data = ImageData("mdv__chart__00000001__", minimal_png_base64)
        assert data.id == "mdv__chart__00000001__"
        assert len(data.png_bytes) > 0
        assert data.width_cm == 15.5

    def test_custom_width(self, minimal_png_base64):
        data = ImageData("mdv__chart__00000001__", minimal_png_base64, width_cm=10.0)
        assert data.width_cm == 10.0

    def test_rejects_oversized_base64(self):
        huge_b64 = "A" * (IMAGE_MAX_B64_LEN + 1)
        with pytest.raises(ValueError, match="exceeds"):
            ImageData("mdv__chart__00000001__", huge_b64)

    def test_rejects_invalid_base64(self):
        with pytest.raises(Exception):
            ImageData("mdv__chart__00000001__", "!!!not-valid-base64!!!")

    def test_rejects_truncated_png(self):
        buf = io.BytesIO()
        img = Image.new("RGB", (10, 10), color=(0, 0, 0))
        img.save(buf, format="PNG")
        full_bytes = buf.getvalue()
        truncated = base64.b64encode(full_bytes[:20]).decode()
        with pytest.raises(Exception):
            ImageData("mdv__chart__00000001__", truncated)

    def test_rejects_pixel_bomb(self):
        buf = io.BytesIO()
        side = 5000
        img = Image.new("RGB", (side, side), color=(0, 0, 0))
        img.save(buf, format="PNG")
        b64 = base64.b64encode(buf.getvalue()).decode()
        if side * side > IMAGE_MAX_PIXELS:
            with pytest.raises(ValueError, match="pixels exceeds"):
                ImageData("mdv__chart__00000001__", b64)


class TestPreprocessMarkdown:
    def test_valid_images(self, minimal_png_base64):
        md = "# Title\n\n![](mdv__chart__aabbccdd__)"
        images = [{"id": "mdv__chart__aabbccdd__", "pngBase64": minimal_png_base64}]
        result_md, image_map = preprocess_markdown(md, images)
        assert result_md == md
        assert "mdv__chart__aabbccdd__" in image_map

    def test_normalizes_alt_text_chart_placeholders(self, minimal_png_base64):
        md = '# Title\n\n![基础流程](mdv__chart__aabbccdd__ "图表")'
        images = [{"id": "mdv__chart__aabbccdd__", "pngBase64": minimal_png_base64}]

        result_md, image_map = preprocess_markdown(md, images)

        assert result_md == "# Title\n\n![](mdv__chart__aabbccdd__)"
        assert "mdv__chart__aabbccdd__" in image_map

    def test_skips_invalid_images(self):
        md = "# Title"
        images = [{"id": "mdv__chart__bad00000__", "pngBase64": "not-valid"}]
        result_md, image_map = preprocess_markdown(md, images)
        assert len(image_map) == 0

    def test_mixed_valid_and_invalid(self, minimal_png_base64):
        md = "text"
        images = [
            {"id": "mdv__chart__good0001__", "pngBase64": minimal_png_base64},
            {"id": "mdv__chart__bad00002__", "pngBase64": "invalid"},
        ]
        _, image_map = preprocess_markdown(md, images)
        assert "mdv__chart__good0001__" in image_map
        assert "mdv__chart__bad00002__" not in image_map

    def test_duplicate_id_last_wins(self, minimal_png_base64, small_png_base64):
        images = [
            {"id": "mdv__chart__dup00001__", "pngBase64": minimal_png_base64},
            {"id": "mdv__chart__dup00001__", "pngBase64": small_png_base64},
        ]
        _, image_map = preprocess_markdown("text", images)
        assert "mdv__chart__dup00001__" in image_map
        assert len(image_map["mdv__chart__dup00001__"].png_bytes) > 100


class TestInjectImages:
    def test_inject_replaces_placeholder(self, tmp_path, small_png_base64):
        placeholder_id = "mdv__chart__a0b1c2d3__"
        md = f"# Test\n\n![](mdv__chart__a0b1c2d3__)\n\nAfter image."

        from app.generator import generate_docx_from_content
        doc_path = str(tmp_path / "inject_test.docx")
        generate_docx_from_content(content=md, output_path=doc_path, style="standard")

        img_data = ImageData(placeholder_id, small_png_base64)
        count = inject_images(doc_path, {placeholder_id: img_data})
        assert count == 1

        doc = Document(doc_path)
        for para in doc.paragraphs:
            assert placeholder_id not in para.text

    def test_empty_image_map_returns_zero(self, tmp_path):
        from app.generator import generate_docx_from_content
        doc_path = str(tmp_path / "empty.docx")
        generate_docx_from_content(content="# Test", output_path=doc_path, style="standard")
        assert inject_images(doc_path, {}) == 0

    def test_preview_image_paragraph_uses_chart_margin(self, tmp_path, small_png_base64):
        placeholder_id = "mdv__chart__a0b1c2d3__"
        md = f"# Test\n\n![]({placeholder_id})\n\nAfter image."

        from app.generator import generate_docx_from_content
        doc_path = str(tmp_path / "preview_image_spacing.docx")
        generate_docx_from_content(content=md, output_path=doc_path, style="preview")

        count = inject_images(
            doc_path,
            {placeholder_id: ImageData(placeholder_id, small_png_base64)},
            style="preview",
        )
        assert count == 1

        doc = Document(doc_path)
        image_para = next(p for p in doc.paragraphs if p._element.xpath(".//w:drawing"))
        assert round(image_para.paragraph_format.space_before.cm, 2) == 0.45
        assert round(image_para.paragraph_format.space_after.cm, 2) == 0.45


class TestImageWidth:
    def test_preview_clamps_image_width_to_content_width(self):
        from app.image_injector import resolve_image_width_cm

        assert resolve_image_width_cm(15.5, style="preview") == 18.5
        assert resolve_image_width_cm(20.0, style="preview") == 19.0
        assert resolve_image_width_cm(15.5, style="standard") == 15.5

    def test_non_preview_image_widths_can_be_clamped_by_layout(self):
        from app.image_injector import resolve_image_width_cm

        assert resolve_image_width_cm(20.0, style="standard", layout=ImageLayout(max_width_cm=15.5)) == 15.5
        assert resolve_image_width_cm(20.0, style="official", layout=ImageLayout(max_width_cm=14.8)) == 14.8

    def test_report_does_not_enlarge_tiny_images(self):
        from app.image_injector import resolve_image_width_cm

        layout = ImageLayout(max_width_cm=15.8, min_width_cm=15.0, min_width_source_threshold_cm=8.0)
        assert resolve_image_width_cm(3.0, style="report", layout=layout) == 3.0
        assert resolve_image_width_cm(10.0, style="report", layout=layout) == 15.0
