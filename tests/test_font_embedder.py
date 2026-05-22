import zipfile

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.font_embedder import (
    DEFAULT_FONT_CANDIDATES,
    DEFAULT_FONT_DIRS,
    embed_fonts_if_requested,
)


def _set_run_font(run, font_name: str):
    run.font.name = font_name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        r_fonts.set(qn(key), font_name)


def test_embed_fonts_without_available_font_keeps_docx_readable(tmp_path):
    doc_path = tmp_path / "sample.docx"
    doc = Document()
    doc.add_paragraph("字体嵌入降级测试")
    doc.save(doc_path)

    warnings = embed_fonts_if_requested(str(doc_path), True, font_paths=["/not/found/font.ttf"])

    reopened = Document(doc_path)
    assert reopened.paragraphs[0].text == "字体嵌入降级测试"
    assert warnings
    assert "未找到可嵌入字体" in warnings[0]
    assert "关闭“嵌入字体”" in warnings[0]
    assert "MD_VIEWER_DOCX_FONT_DIRS" in warnings[0]


def test_embed_fonts_disabled_returns_no_warnings(tmp_path):
    doc_path = tmp_path / "sample.docx"
    doc = Document()
    doc.add_paragraph("no-op")
    doc.save(doc_path)

    assert embed_fonts_if_requested(str(doc_path), False) == []


def test_default_font_policy_does_not_scan_macos_system_font_dirs():
    default_locations = [*(str(path) for path in DEFAULT_FONT_DIRS), *DEFAULT_FONT_CANDIDATES]

    assert not any(location.startswith("/System/Library/Fonts") for location in default_locations)
    assert not any(location.startswith("/Library/Fonts") for location in default_locations)
    assert not any(location.endswith("/Library/Fonts") for location in default_locations)


def test_embed_fonts_uses_env_font_paths_by_default(tmp_path, monkeypatch):
    font_path = tmp_path / "CustomCjk.ttf"
    font_path.write_bytes(b"fake-font-bytes")
    monkeypatch.setenv("MD_VIEWER_DOCX_FONT_PATHS", str(font_path))

    doc_path = tmp_path / "sample.docx"
    doc = Document()
    run = doc.add_paragraph().add_run("env font")
    _set_run_font(run, "CustomCjk")
    doc.save(doc_path)

    warnings = embed_fonts_if_requested(str(doc_path), True)

    assert warnings == []
    import zipfile
    with zipfile.ZipFile(doc_path) as zf:
        assert "word/fonts/CustomCjk.ttf" in zf.namelist()


def test_embed_fonts_writes_all_referenced_font_files(tmp_path):
    first_font = tmp_path / "FirstCjk.ttf"
    second_font = tmp_path / "SecondCjk.otf"
    first_font.write_bytes(b"first-font-bytes")
    second_font.write_bytes(b"second-font-bytes")

    doc_path = tmp_path / "multi-font.docx"
    doc = Document()
    _set_run_font(doc.add_paragraph().add_run("first font"), "FirstCjk")
    _set_run_font(doc.add_paragraph().add_run("second font"), "SecondCjk")
    doc.save(doc_path)

    warnings = embed_fonts_if_requested(
        str(doc_path),
        True,
        font_paths=[str(first_font), str(second_font)],
    )

    assert warnings == []
    with zipfile.ZipFile(doc_path) as zf:
        names = zf.namelist()
    assert "word/fonts/FirstCjk.ttf" in names
    assert "word/fonts/SecondCjk.otf" in names


def test_embed_fonts_only_writes_docx_referenced_font_files(tmp_path):
    first_font = tmp_path / "FirstCjk.ttf"
    second_font = tmp_path / "SecondCjk.otf"
    first_font.write_bytes(b"first-font-bytes")
    second_font.write_bytes(b"second-font-bytes")

    doc_path = tmp_path / "referenced-font.docx"
    doc = Document()
    run = doc.add_paragraph().add_run("only first font is referenced")
    _set_run_font(run, "FirstCjk")
    doc.save(doc_path)

    warnings = embed_fonts_if_requested(
        str(doc_path),
        True,
        font_paths=[str(first_font), str(second_font)],
    )

    assert warnings == []
    with zipfile.ZipFile(doc_path) as zf:
        names = zf.namelist()
    assert "word/fonts/FirstCjk.ttf" in names
    assert "word/fonts/SecondCjk.otf" not in names


def test_embed_fonts_matches_noto_cjk_ttc_regional_face(tmp_path):
    font_path = tmp_path / "NotoSansCJK-Regular.ttc"
    font_path.write_bytes(b"noto-cjk-ttc-bytes")

    doc_path = tmp_path / "noto-face.docx"
    doc = Document()
    run = doc.add_paragraph().add_run("docker noto font")
    _set_run_font(run, "Noto Sans CJK SC")
    doc.save(doc_path)

    warnings = embed_fonts_if_requested(
        str(doc_path),
        True,
        font_paths=[str(font_path)],
    )

    assert warnings == []
    with zipfile.ZipFile(doc_path) as zf:
        names = zf.namelist()
    assert "word/fonts/NotoSansCJK-Regular.ttc" in names


def test_embed_fonts_warning_explains_how_to_fix_unmatched_fonts(tmp_path):
    font_path = tmp_path / "NotoSansCJKsc-Regular.otf"
    font_path.write_bytes(b"noto-cjk-bytes")

    doc_path = tmp_path / "unmatched-font.docx"
    doc = Document()
    run = doc.add_paragraph().add_run("uses macOS font")
    _set_run_font(run, "PingFang SC")
    doc.save(doc_path)

    warnings = embed_fonts_if_requested(
        str(doc_path),
        True,
        font_paths=[str(font_path)],
    )

    assert len(warnings) == 1
    assert "PingFang SC" in warnings[0]
    assert "关闭“嵌入字体”" in warnings[0]
    assert "服务端挂载授权字体" in warnings[0]
    assert "MD_VIEWER_DOCX_FONT_PATHS" in warnings[0]
