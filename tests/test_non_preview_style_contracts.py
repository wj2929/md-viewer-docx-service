import json
import io
import os
import shutil
import subprocess
import zipfile
from pathlib import Path
from unittest.mock import patch

import pytest
from docx import Document
from lxml import etree

from app.image_injector import ImageData, inject_images
from app.generator import generate_docx_from_content


NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
NON_PREVIEW_STYLES = ("standard", "official", "internal", "report")
FORMAL_STYLES = ("official", "internal", "report")
FIXTURE_DIR = Path(__file__).parent / "fixtures" / "non_preview_styles"
FIXTURE_NAMES = (
    "simple.md",
    "headings-h1-h6.md",
    "nested-lists.md",
    "official-document.md",
    "internal-notice.md",
    "report-with-charts.md",
    "wide-table.md",
    "image-layout.md",
)
TABLE_FIXTURES = {
    "simple.md",
    "internal-notice.md",
    "report-with-charts.md",
    "wide-table.md",
    "image-layout.md",
}


def _xml_tree(path):
    with zipfile.ZipFile(path) as zf:
        return etree.parse(zf.open("word/document.xml"))


def _word_xml_text(path):
    with zipfile.ZipFile(path) as zf:
        return "\n".join(
            zf.read(name).decode("utf-8", errors="ignore")
            for name in zf.namelist()
            if name.startswith("word/") and name.endswith(".xml")
        )


def _fills(path):
    return _xml_tree(path).xpath("//w:shd/@w:fill", namespaces=NS)


def _table_margins(path):
    return _xml_tree(path).xpath("//w:tcMar", namespaces=NS)


def _left_borders(path):
    return _xml_tree(path).xpath("//w:pBdr/w:left", namespaces=NS)


def _paragraph_by_text(path, text):
    matches = _xml_tree(path).xpath(
        "//w:p[.//w:t[text()=$text]]",
        namespaces=NS,
        text=text,
    )
    assert matches, f"paragraph not found: {text}"
    return matches[0]


def _fixture_text(name):
    return (FIXTURE_DIR / name).read_text(encoding="utf-8")


def test_non_preview_fixture_matrix_is_available():
    assert sorted(path.name for path in FIXTURE_DIR.glob("*.md")) == sorted(FIXTURE_NAMES)
    for name in FIXTURE_NAMES:
        text = _fixture_text(name)
        assert text.startswith("# "), name
        assert len(text.strip()) > 20, name


@pytest.mark.parametrize("style", NON_PREVIEW_STYLES)
@pytest.mark.parametrize("fixture_name", FIXTURE_NAMES)
def test_non_preview_fixture_matrix_generates_structural_docx(tmp_path, style, fixture_name):
    out_path = tmp_path / f"{style}-{fixture_name.replace('.md', '.docx')}"

    generate_docx_from_content(
        content=_fixture_text(fixture_name),
        output_path=str(out_path),
        style=style,
    )

    with zipfile.ZipFile(out_path) as zf:
        names = set(zf.namelist())
    assert "word/document.xml" in names
    assert "word/styles.xml" in names

    doc = Document(out_path)
    assert doc.paragraphs or doc.tables
    assert round(doc.sections[0].page_width.cm, 1) == 21.0
    assert round(doc.sections[0].page_height.cm, 1) == 29.7

    xml = _word_xml_text(out_path)
    if style in FORMAL_STYLES:
        assert "由 MD Viewer 生成" not in xml

    if fixture_name in TABLE_FIXTURES:
        tree = _xml_tree(out_path)
        assert tree.xpath("//w:tbl", namespaces=NS), fixture_name
        for row in tree.xpath("//w:tbl/w:tr[1]", namespaces=NS):
            assert row.xpath("./w:trPr/w:tblHeader", namespaces=NS), fixture_name


def test_official_fixture_preserves_h5_heading_and_list_numbering(tmp_path):
    out_path = tmp_path / "official-document.docx"

    generate_docx_from_content(
        content=_fixture_text("official-document.md"),
        output_path=str(out_path),
        style="official",
    )

    xml = _word_xml_text(out_path)
    assert 'w:pStyle w:val="Heading5"' in xml
    assert 'w:eastAsia="仿宋_GB2312"' in xml

    list_out = tmp_path / "official-nested-list.docx"
    generate_docx_from_content(
        content=_fixture_text("nested-lists.md"),
        output_path=str(list_out),
        style="official",
    )
    texts = [p.text.strip() for p in Document(list_out).paragraphs]
    assert "1.  第一项" in texts
    assert "1.  第一项子项" in texts
    assert "2.  第二项" in texts


def test_non_preview_fixture_image_placeholders_are_injected(tmp_path, small_png_base64):
    placeholder_id = "mdv__chart__a0b1c2d3__"
    out_path = tmp_path / "image-layout.docx"

    generate_docx_from_content(
        content=_fixture_text("image-layout.md"),
        output_path=str(out_path),
        style="report",
    )
    injected = inject_images(
        str(out_path),
        {placeholder_id: ImageData(placeholder_id, small_png_base64)},
        style="report",
    )

    assert injected == 2
    xml = _word_xml_text(out_path)
    assert placeholder_id not in xml
    assert "<w:drawing>" in xml


def test_convert_warning_header_uses_non_preview_fixture(client):
    with patch("app.main._font_status_by_style", return_value={
        "official": {
            "仿宋_GB2312": {
                "status": "fallback",
                "resolved": "Noto Sans CJK SC",
                "fallback": "Noto Sans CJK SC",
                "embeddable": True,
            }
        }
    }):
        resp = client.post("/convert", json={
            "markdown": _fixture_text("official-document.md"),
            "style": "official",
        })

    assert resp.status_code == 200
    warnings = json.loads(resp.headers["x-convert-warnings"])
    assert any("仿宋_GB2312" in warning and "Noto Sans CJK SC" in warning for warning in warnings)
    assert "由 MD Viewer 生成" not in _word_xml_text_from_bytes(resp.content)


def _word_xml_text_from_bytes(docx_bytes):
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zf:
        return "\n".join(
            zf.read(name).decode("utf-8", errors="ignore")
            for name in zf.namelist()
            if name.startswith("word/") and name.endswith(".xml")
        )


def test_non_preview_official_fixture_converts_to_a4_pdf_when_tools_available(tmp_path):
    soffice = shutil.which("soffice") or (
        "/Applications/LibreOffice.app/Contents/MacOS/soffice"
        if os.path.exists("/Applications/LibreOffice.app/Contents/MacOS/soffice")
        else None
    )
    pdfinfo = shutil.which("pdfinfo")
    if not soffice or not pdfinfo:
        pytest.skip("soffice/pdfinfo not available")

    out_path = tmp_path / "official-document.docx"
    generate_docx_from_content(
        content=_fixture_text("official-document.md"),
        output_path=str(out_path),
        style="official",
    )

    subprocess.run(
        [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(tmp_path), str(out_path)],
        check=True,
        capture_output=True,
        text=True,
        timeout=60,
    )
    result = subprocess.run(
        [pdfinfo, str(tmp_path / "official-document.pdf")],
        check=True,
        capture_output=True,
        text=True,
        timeout=20,
    )

    assert "Pages:" in result.stdout
    assert "A4" in result.stdout or "595." in result.stdout


def test_non_preview_styles_use_a4_page_size(tmp_path):
    for style in ("standard", "official", "internal", "report"):
        out_path = tmp_path / f"{style}-a4.docx"
        generate_docx_from_content(
            content="# 标题\n\n正文",
            output_path=str(out_path),
            style=style,
        )
        section = Document(out_path).sections[0]
        assert round(section.page_width.cm, 1) == 21.0
        assert round(section.page_height.cm, 1) == 29.7


def test_non_preview_headings_disable_word_keep_constraints(tmp_path):
    for style in ("standard", "official", "internal", "report"):
        out_path = tmp_path / f"{style}-heading-flow.docx"
        generate_docx_from_content(
            content="# 标题\n\n> 摘要说明\n\n## 1. 图表\n\n### 1.1 流程图\n\n![](mdv__chart__deadbeef__)",
            output_path=str(out_path),
            style=style,
        )

        for heading_text in ("1. 图表", "1.1 流程图"):
            para = _paragraph_by_text(out_path, heading_text)
            keep_next = para.xpath("./w:pPr/w:keepNext/@w:val", namespaces=NS)
            keep_lines = para.xpath("./w:pPr/w:keepLines/@w:val", namespaces=NS)
            assert keep_next == ["0"]
            assert keep_lines == ["0"]


def test_document_body_first_line_indent_matches_style_contract(tmp_path):
    expected_indent_pt = {
        "preview": None,
        "standard": None,
        "official": 32.0,
        "internal": 30.0,
        "report": 24.0,
    }
    for style, expected_pt in expected_indent_pt.items():
        out_path = tmp_path / f"{style}-body-indent.docx"
        generate_docx_from_content(
            content="# 标题\n\n这是一段普通正文，用于检查不同 DOCX 格式的首行缩进契约。",
            output_path=str(out_path),
            style=style,
        )

        para = next(p for p in Document(out_path).paragraphs if p.text.startswith("这是一段普通正文"))
        if expected_pt is None:
            assert para.paragraph_format.first_line_indent is None
        else:
            assert round(para.paragraph_format.first_line_indent.pt, 2) == expected_pt


def test_non_preview_image_layout_caps_page_height():
    from app.main import _image_layout_for_style

    for style in ("standard", "official", "internal", "report"):
        layout = _image_layout_for_style(style)
        assert layout is not None
        assert layout.max_height_cm <= 15.0


def test_standard_table_has_header_fill_and_cell_margins(tmp_path):
    out_path = tmp_path / "standard-table.docx"
    generate_docx_from_content(
        content="# 标题\n\n| A | B |\n|---|---|\n| 1 | 2 |",
        output_path=str(out_path),
        style="standard",
    )

    assert "F6F8FA" in _fills(out_path)
    assert _table_margins(out_path)

    doc = Document(out_path)
    header_run = next(run for run in doc.tables[0].cell(0, 0).paragraphs[0].runs if run.text)
    body_run = next(run for run in doc.tables[0].cell(1, 0).paragraphs[0].runs if run.text)
    assert header_run.font.size.pt == 9.5
    assert body_run.font.size.pt == 9.5


def test_non_preview_table_header_repeats_across_pages(tmp_path):
    for style in ("standard", "official", "internal", "report"):
        out_path = tmp_path / f"{style}-repeat-header.docx"
        generate_docx_from_content(
            content="# 标题\n\n| A | B |\n|---|---|\n| 1 | 2 |\n| 3 | 4 |",
            output_path=str(out_path),
            style=style,
        )

        first_row = _xml_tree(out_path).xpath("//w:tbl[1]/w:tr[1]", namespaces=NS)[0]
        assert first_row.xpath("./w:trPr/w:tblHeader", namespaces=NS), style


def test_official_table_does_not_use_web_header_fill(tmp_path):
    out_path = tmp_path / "official-table.docx"
    generate_docx_from_content(
        content="# 标题\n\n| A | B |\n|---|---|\n| 1 | 2 |",
        output_path=str(out_path),
        style="official",
    )

    assert "F6F8FA" not in _fills(out_path)
    assert _table_margins(out_path)


def test_official_note_prefix_only_for_explicit_note(tmp_path):
    out_path = tmp_path / "official-note.docx"
    generate_docx_from_content(
        content="# 标题\n\n> **注意：** 这是正式公文里的说明。",
        output_path=str(out_path),
        style="official",
    )

    doc = Document(out_path)
    assert any(p.text.startswith("注：") for p in doc.paragraphs)
    assert not _left_borders(out_path)


def test_official_normal_quote_does_not_become_note(tmp_path):
    out_path = tmp_path / "official-quote.docx"
    generate_docx_from_content(
        content="# 标题\n\n> 引用一段政策原文。",
        output_path=str(out_path),
        style="official",
    )

    doc = Document(out_path)
    assert not any(p.text.startswith("注：") for p in doc.paragraphs)
    assert any("引用一段政策原文" in p.text for p in doc.paragraphs)


def test_official_quote_uses_first_line_indent_not_left_indent(tmp_path):
    out_path = tmp_path / "official-quote-indent.docx"
    generate_docx_from_content(
        content="# 标题\n\n> 来源：这是一段比较长的说明文字，用来模拟导出后可能换行的家长会来源信息。",
        output_path=str(out_path),
        style="official",
    )

    para = next(p for p in Document(out_path).paragraphs if p.text.startswith("来源："))
    assert para.paragraph_format.left_indent is None
    assert round(para.paragraph_format.first_line_indent.pt, 2) == 32.0


def test_official_quote_lead_in_before_list_uses_first_line_indent(tmp_path):
    out_path = tmp_path / "official-quote-list-lead-in-indent.docx"
    generate_docx_from_content(
        content="# 标题\n\n> **结论**：如需按分部/频道分析连麦使用情况，只能通过以下方式：\n> 1. 第一种方式\n> 2. 第二种方式",
        output_path=str(out_path),
        style="official",
    )

    para = next(p for p in Document(out_path).paragraphs if p.text.startswith("结论："))
    assert para.paragraph_format.left_indent is None
    assert round(para.paragraph_format.first_line_indent.pt, 2) == 32.0


def test_official_quote_list_items_use_hanging_indent(tmp_path):
    out_path = tmp_path / "official-quote-list-indent.docx"
    generate_docx_from_content(
        content="# 标题\n\n> **说明**：\n> - 第一项说明\n> - 第二项说明",
        output_path=str(out_path),
        style="official",
    )

    doc = Document(out_path)
    first_item = next(p for p in doc.paragraphs if "第一项说明" in p.text)
    assert not first_item.text.strip().startswith("-")
    assert first_item.text.strip().startswith("•\t")
    assert round(first_item.paragraph_format.left_indent.cm, 2) == 1.2
    assert round(first_item.paragraph_format.first_line_indent.cm, 2) == -0.5
    assert [round(tab.position.cm, 2) for tab in first_item.paragraph_format.tab_stops] == [1.2]


def test_official_quote_ordered_list_items_use_hanging_indent(tmp_path):
    out_path = tmp_path / "official-quote-ordered-list-indent.docx"
    generate_docx_from_content(
        content="# 标题\n\n> **建议**：\n> 1. 第一项建议\n> 2. 第二项建议",
        output_path=str(out_path),
        style="official",
    )

    doc = Document(out_path)
    first_item = next(p for p in doc.paragraphs if "第一项建议" in p.text)
    second_item = next(p for p in doc.paragraphs if "第二项建议" in p.text)
    assert first_item.text.strip().startswith("1.\t")
    assert second_item.text.strip().startswith("2.\t")
    assert round(first_item.paragraph_format.left_indent.cm, 2) == 1.2
    assert round(first_item.paragraph_format.first_line_indent.cm, 2) == -0.5
    assert [round(tab.position.cm, 2) for tab in first_item.paragraph_format.tab_stops] == [1.2]


def test_official_gfm_note_becomes_note_prefix(tmp_path):
    out_path = tmp_path / "official-gfm-note.docx"
    generate_docx_from_content(
        content="# 标题\n\n> [!NOTE]\n> 这是提示。",
        output_path=str(out_path),
        style="official",
    )

    doc = Document(out_path)
    assert any(p.text.startswith("注：") and "这是提示" in p.text for p in doc.paragraphs)


def test_report_note_preserves_inline_markdown(tmp_path):
    out_path = tmp_path / "report-note-inline.docx"
    generate_docx_from_content(
        content="# 标题\n\n> **注意：** 使用 `kubectl` 检查。",
        output_path=str(out_path),
        style="report",
    )

    doc = Document(out_path)
    assert doc.tables
    assert "kubectl" in doc.tables[0].cell(0, 0).text


def test_official_code_block_uses_9pt_single_spacing(tmp_path):
    out_path = tmp_path / "official-code.docx"
    generate_docx_from_content(
        content="# 标题\n\n```bash\nkubectl get pvc\n```",
        output_path=str(out_path),
        style="official",
    )

    doc = Document(out_path)
    para = next(p for p in doc.paragraphs if "kubectl get pvc" in p.text)
    run = next(r for r in para.runs if r.text)
    assert run.font.size.pt == 9.0
    assert para.paragraph_format.line_spacing == 1.0
