import os
import pytest
import zipfile
from docx import Document

from app.generator import parse_inline, parse_markdown, generate_docx_from_content, BlockType


class TestParseInline:
    def test_bold(self):
        runs = parse_inline("**粗体** text")
        bold_runs = [r for r in runs if r.bold]
        assert len(bold_runs) == 1
        assert bold_runs[0].text == "粗体"

    def test_bold_inside_chinese_sentence(self):
        runs = parse_inline("父母在京**务工就业证明**")
        assert [(r.text, r.bold) for r in runs] == [
            ("父母在京", False),
            ("务工就业证明", True),
        ]

    def test_italic(self):
        runs = parse_inline("*斜体* text")
        italic_runs = [r for r in runs if r.italic]
        assert len(italic_runs) == 1
        assert italic_runs[0].text == "斜体"

    def test_code(self):
        runs = parse_inline("使用 `print()` 函数")
        code_runs = [r for r in runs if r.code]
        assert len(code_runs) == 1
        assert code_runs[0].text == "print()"

    def test_link(self):
        runs = parse_inline("点击 [链接](https://example.com) 跳转")
        link_runs = [r for r in runs if r.link]
        assert len(link_runs) == 1
        assert link_runs[0].text == "链接"
        assert link_runs[0].link == "https://example.com"

    def test_plain_text(self):
        runs = parse_inline("纯文本内容")
        assert len(runs) == 1
        assert runs[0].text == "纯文本内容"
        assert not runs[0].bold and not runs[0].italic and not runs[0].code

    def test_decodes_html_entities_in_text_runs(self):
        runs = parse_inline("统计周期&nbsp;&nbsp;|&nbsp;&nbsp;数据来源 &amp; 说明")

        assert "".join(run.text for run in runs) == "统计周期  |  数据来源 & 说明"

    def test_keeps_html_entities_literal_inside_code_spans(self):
        runs = parse_inline("使用 `&nbsp;` 表示空格")

        assert [(run.text, run.code) for run in runs] == [
            ("使用 ", False),
            ("&nbsp;", True),
            (" 表示空格", False),
        ]


class TestParseMarkdown:
    def test_heading(self):
        blocks = parse_markdown("# 标题\n\n正文")
        assert blocks[0].type == BlockType.HEADING
        assert blocks[0].level == 1

    def test_paragraph(self):
        blocks = parse_markdown("这是一段正文。")
        assert any(b.type == BlockType.PARAGRAPH for b in blocks)

    def test_unordered_list(self):
        blocks = parse_markdown("- 项目 1\n- 项目 2")
        assert any(b.type == BlockType.UNORDERED_LIST for b in blocks)

    def test_ordered_list(self):
        blocks = parse_markdown("1. 第一项\n2. 第二项")
        assert any(b.type == BlockType.ORDERED_LIST for b in blocks)

    def test_table(self):
        md = "| A | B |\n|---|---|\n| 1 | 2 |"
        blocks = parse_markdown(md)
        assert any(b.type == BlockType.TABLE for b in blocks)

    def test_code_block(self):
        md = "```python\nprint('hello')\n```"
        blocks = parse_markdown(md)
        assert any(b.type == BlockType.CODE_BLOCK for b in blocks)

    def test_blockquote(self):
        blocks = parse_markdown("> 引用文本")
        assert any(b.type == BlockType.BLOCKQUOTE for b in blocks)

    def test_horizontal_rule(self):
        blocks = parse_markdown("---")
        assert any(b.type == BlockType.HORIZONTAL_RULE for b in blocks)


class TestGenerateDocx:
    @pytest.mark.parametrize("style", ["preview", "standard", "official", "internal", "report"])
    def test_generate_each_style(self, tmp_path, sample_markdown, style):
        out_path = str(tmp_path / f"out_{style}.docx")
        generate_docx_from_content(
            content=sample_markdown,
            output_path=out_path,
            style=style,
        )
        assert os.path.exists(out_path)
        assert os.path.getsize(out_path) > 0

    @pytest.mark.parametrize("style", ["preview", "standard", "official", "internal", "report"])
    def test_none_footer_text_omits_generated_branding(self, tmp_path, style):
        out_path = str(tmp_path / f"no-footer-{style}.docx")
        generate_docx_from_content(
            content="# 标题\n\n正文",
            output_path=out_path,
            style=style,
            footer_text=None,
        )

        with zipfile.ZipFile(out_path) as zf:
            xml = zf.read("word/document.xml").decode("utf-8")

        assert "由 MD Viewer 生成" not in xml

    def test_preview_uses_a4_and_narrow_margins(self, tmp_path):
        out_path = str(tmp_path / "preview.docx")
        generate_docx_from_content(
            content="# 标题\n\n## 小节\n\n正文",
            output_path=out_path,
            style="preview",
        )

        doc = Document(out_path)
        section = doc.sections[0]
        assert round(section.page_width.cm, 1) == 21.0
        assert round(section.page_height.cm, 1) == 29.7
        assert section.left_margin.cm <= 1.3
        assert section.right_margin.cm <= 1.3
        assert section.top_margin.cm <= 1.3
        assert section.bottom_margin.cm <= 1.3

    def test_preview_font_fallback_returns_non_empty_font(self):
        from app.generator import _resolve_preview_fonts

        body_font, mono_font = _resolve_preview_fonts()
        assert isinstance(body_font, str)
        assert isinstance(mono_font, str)
        assert body_font
        assert mono_font

    def test_preview_table_has_header_shading(self, tmp_path):
        out_path = str(tmp_path / "preview-table.docx")
        generate_docx_from_content(
            content="| A | B | C |\n|---|---|---|\n| 1 | 2 | 3 |",
            output_path=out_path,
            style="preview",
        )

        with zipfile.ZipFile(out_path) as zf:
            xml = zf.read("word/document.xml").decode("utf-8")
        assert 'w:fill="F6F8FA"' in xml
        assert "w:tcMar" in xml

    def test_docx_output_decodes_html_entities_in_blockquote(self, tmp_path):
        out_path = str(tmp_path / "html-entities.docx")
        generate_docx_from_content(
            content="> **统计周期**：2021年3月 — 2026年5月 &nbsp;&nbsp;|&nbsp;&nbsp; **数据来源**：开放 API",
            output_path=out_path,
            style="official",
        )

        with zipfile.ZipFile(out_path) as zf:
            xml = zf.read("word/document.xml").decode("utf-8")
        assert "&amp;nbsp;" not in xml
        assert "2026年5月   |   " in xml

    def test_preview_table_cell_metrics_match_print_density(self, tmp_path):
        out_path = str(tmp_path / "preview-table-metrics.docx")
        generate_docx_from_content(
            content="| 存储类（StorageClass） | 华为云服务 |\n|---|---|\n| `csi-disk` | EVS（云硬盘） |",
            output_path=out_path,
            style="preview",
        )

        doc = Document(out_path)
        header_run = next(run for run in doc.tables[0].cell(0, 0).paragraphs[0].runs if run.text)
        data_run = next(run for run in doc.tables[0].cell(1, 1).paragraphs[0].runs if run.text)
        assert header_run.font.size.pt == 8.5
        assert data_run.font.size.pt == 8.0

        with zipfile.ZipFile(out_path) as zf:
            xml = zf.read("word/document.xml").decode("utf-8")
        assert 'w:top w:w="55"' in xml
        assert 'w:start w:w="110"' in xml
        assert 'w:bottom w:w="55"' in xml
        assert 'w:end w:w="110"' in xml

    def test_preview_two_column_table_uses_content_width(self):
        from app.generator import _preview_table_target_width_cm

        width = _preview_table_target_width_cm([
            ["属性", "值"],
            ["PVC", "cce-obs-ilesson"],
            ["存储类型", "OBS 对象存储（csi-obs），obsfs 挂载"],
            ["使用者", "dash-worker (12+10+10=32 pods), wrpc-office (3+4+6=13 pods)"],
            ["数据流", "录制/上传的原始课件 → OBS 桶 → dash-worker 读取处理 → 输出到 outbound"],
        ])

        assert 14.5 <= width <= 16.5

    def test_preview_wide_table_still_uses_content_width(self):
        from app.generator import _preview_table_target_width_cm

        width = _preview_table_target_width_cm([
            ["存储类（StorageClass）", "华为云服务", "存储类型说明", "PVC 数量", "Bound", "Pending", "总容量"],
            ["`csi-disk`", "EVS（云硬盘）", "普通 IO 云硬盘", "5", "5", "0", "248 Gi"],
            ["`csi-disk-ssd`", "EVS（云硬盘）", "超高 IO 云硬盘（ESSD）", "3", "3", "0", "370 Gi"],
            ["`csi-disk-topology`", "EVS（云硬盘）", "拓扑感知云硬盘（SAS）", "3", "3", "0", "205 Gi"],
        ])

        assert round(width, 1) == 19.0

    def test_preview_table_has_margin_after_table(self, tmp_path):
        out_path = str(tmp_path / "preview-table-gap.docx")
        generate_docx_from_content(
            content="| A | B |\n|---|---|\n| 1 | 2 |\n\n表格后的正文",
            output_path=out_path,
            style="preview",
        )

        with zipfile.ZipFile(out_path) as zf:
            xml = zf.read("word/document.xml").decode("utf-8")
        assert '<w:tbl>' in xml
        assert 'w:after="80"' in xml

    def test_preview_heading_spacing_is_compact(self, tmp_path):
        out_path = str(tmp_path / "preview-heading-spacing.docx")
        generate_docx_from_content(
            content="# 标题\n\n## 一、存储总览\n\n正文",
            output_path=out_path,
            style="preview",
        )

        doc = Document(out_path)
        heading = next(p for p in doc.paragraphs if p.text == "一、存储总览")
        assert heading.paragraph_format.space_before.pt == 8
        assert heading.paragraph_format.space_after.pt == 4

    def test_preview_heading_is_not_italic(self, tmp_path):
        out_path = str(tmp_path / "preview-heading.docx")
        generate_docx_from_content(
            content="# 标题\n\n#### /var/lib/mysql — WRPC 数据库",
            output_path=out_path,
            style="preview",
        )

        with zipfile.ZipFile(out_path) as zf:
            xml = zf.read("word/document.xml").decode("utf-8")
        assert 'w:i w:val="0"' in xml

    def test_preview_uses_resolved_font_consistently(self, tmp_path):
        from app.generator import _resolve_preview_fonts

        body_font, _ = _resolve_preview_fonts()
        out_path = str(tmp_path / "preview-fonts.docx")
        generate_docx_from_content(
            content="# 标题\n\n正文 English",
            output_path=out_path,
            style="preview",
        )

        with zipfile.ZipFile(out_path) as zf:
            xml = zf.read("word/document.xml").decode("utf-8")
        assert f'w:ascii="{body_font}"' in xml
        assert f'w:hAnsi="{body_font}"' in xml
        assert f'w:eastAsia="{body_font}"' in xml

    def test_preview_list_indent_and_links_match_browser_preview(self, tmp_path):
        out_path = str(tmp_path / "preview-list-link.docx")
        generate_docx_from_content(
            content="# 标题\n\n- [一、存储总览](#一存储总览)\n- [二、系统全景](#二系统全景)",
            output_path=out_path,
            style="preview",
        )

        doc = Document(out_path)
        first_item = next(p for p in doc.paragraphs if "一、存储总览" in p.text)
        assert round(first_item.paragraph_format.left_indent.cm, 2) == 0.85
        assert round(first_item.paragraph_format.first_line_indent.cm, 2) == -0.32
        assert first_item.paragraph_format.space_before.pt == 0
        assert first_item.paragraph_format.space_after.pt == 0
        assert first_item.paragraph_format.line_spacing == 1.25

        with zipfile.ZipFile(out_path) as zf:
            xml = zf.read("word/document.xml").decode("utf-8")
        assert 'w:u w:val="none"' in xml
        assert 'w:u w:val="single"' not in xml

    def test_official_lists_preserve_inline_bold_without_markers(self, tmp_path):
        out_path = str(tmp_path / "official-list-bold.docx")
        generate_docx_from_content(
            content="# 标题\n\n1. 父母在京**务工就业证明**\n\n- 可直接进**一中实验班/清北班**",
            output_path=out_path,
            style="official",
        )

        doc = Document(out_path)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "**" not in text
        assert any(run.text == "务工就业证明" and run.bold for p in doc.paragraphs for run in p.runs)
        assert any(run.text == "一中实验班/清北班" and run.bold for p in doc.paragraphs for run in p.runs)

    def test_preview_preserves_horizontal_rule_before_heading(self, tmp_path):
        out_path = str(tmp_path / "preview-horizontal-rule.docx")
        generate_docx_from_content(
            content="# 标题\n\n- 目录项\n\n---\n\n## 一、存储总览",
            output_path=out_path,
            style="preview",
        )

        with zipfile.ZipFile(out_path) as zf:
            xml = zf.read("word/document.xml").decode("utf-8")
        assert "<w:pBdr>" in xml
        assert 'w:bottom w:val="single"' in xml
        assert 'w:color="EAECEF"' in xml

    def test_preview_code_block_has_shading(self, tmp_path):
        out_path = str(tmp_path / "preview-code.docx")
        generate_docx_from_content(
            content="```bash\nkubectl get pvc --all-namespaces\n```",
            output_path=out_path,
            style="preview",
        )

        with zipfile.ZipFile(out_path) as zf:
            xml = zf.read("word/document.xml").decode("utf-8")
        assert 'w:fill="F6F8FA"' in xml
        doc = Document(out_path)
        assert len(doc.paragraphs) > 0

    def test_preview_note_block_is_callout_without_quote_bar(self, tmp_path):
        out_path = str(tmp_path / "preview-note.docx")
        generate_docx_from_content(
            content="> **注意：** Deployment 中 volume 名称仍叫 cce-sfs-ilesson，但实际绑定的 PVC 已经是 cce-obs-ilesson。",
            output_path=out_path,
            style="preview",
        )

        doc = Document(out_path)
        assert len(doc.tables) == 1
        note_cell = doc.tables[0].cell(0, 0)
        assert note_cell.text.startswith("注意：")
        first_run = next(run for run in note_cell.paragraphs[0].runs if run.text)
        assert first_run.font.size.pt == 9.5

        with zipfile.ZipFile(out_path) as zf:
            xml = zf.read("word/document.xml").decode("utf-8")
        assert 'w:fill="F6F8FA"' in xml
        assert '<w:pBdr>' not in xml
        assert 'w:left w:val="single"' not in xml

    def test_preview_table_column_widths_follow_content_weight(self):
        from app.generator import _preview_column_widths_cm

        widths = _preview_column_widths_cm([
            ["存储类（StorageClass）", "华为云服务", "存储类型说明", "PVC 数量", "Bound", "Pending", "总容量"],
            ["`csi-disk`", "EVS（云硬盘）", "普通 IO 云硬盘", "5", "5", "0", "248 Gi"],
            ["`csi-disk-ssd`", "EVS（云硬盘）", "超高 IO 云硬盘（ESSD）", "3", "3", "0", "370 Gi"],
            ["`csi-disk-topology`", "EVS（云硬盘）", "拓扑感知云硬盘（SAS）", "3", "3", "0", "205 Gi"],
            ["`csi-nas`", "SFS（弹性文件服务）", "NFS 共享文件存储", "4", "4", "0", "1520 Gi"],
            ["`csi-obs`", "OBS（对象存储）", "对象存储桶挂载", "2", "2", "0", "2 Gi"],
            ["`csi-sfs`", "SFS（弹性文件服务）", "SFS 文件共享", "1", "1", "0", "1 Gi"],
            ["`csi-sfsturbo`", "SFS Turbo（极速文件存储）", "高性能共享文件存储", "2", "1", "1", "500 Gi"],
            ["**合计**", "", "", "**20**", "**19**", "**1**", "**2846 Gi**"],
        ])

        assert len(widths) == 7
        assert round(sum(widths), 1) == 19.0
        assert widths[0] > widths[3]
        assert widths[2] > widths[4]
        assert widths[0] >= 4.1
        assert widths[1] >= 4.3
        assert widths[2] >= 3.9
        assert widths[3] >= 1.6
        assert all(width <= 1.7 for width in widths[3:])

    def test_auto_extract_title(self, tmp_path):
        md = "# 自动提取的标题\n\n正文内容"
        out_path = str(tmp_path / "title_test.docx")
        generate_docx_from_content(content=md, output_path=out_path, style="standard")
        doc = Document(out_path)
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert any("自动提取的标题" in t for t in texts)

    def test_invalid_style_falls_back_to_standard(self, tmp_path):
        out_path = str(tmp_path / "fallback.docx")
        generate_docx_from_content(
            content="# Test\n\nBody",
            output_path=out_path,
            style="nonexistent",
        )
        assert os.path.exists(out_path)
        assert os.path.getsize(out_path) > 0
