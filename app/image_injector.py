"""
图片占位符注入模块

将 markdown 中的 ![](mdv__chart__xxxxxxxx__) 占位符替换为实际图片。
在 python-docx 生成 DOCX 后，扫描段落找到占位符文本，替换为 InlineImage。
"""
import re
import base64
import io
import logging
from dataclasses import dataclass
from typing import Dict, Optional
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from PIL import Image

logger = logging.getLogger(__name__)

PLACEHOLDER_PATTERN = re.compile(r"^!\[\]\((mdv__chart__[0-9a-f]{8}__)\)$")
PLACEHOLDER_IMAGE_MARKDOWN_PATTERN = re.compile(
    r"!\[[^\]\n]*\]\(\s*(mdv__chart__[0-9a-f]{8}__)(?:\s+(?:\"[^\"]*\"|'[^']*'|\([^)]*\)))?\s*\)"
)
IMAGE_MAX_PIXELS = 20_000_000
IMAGE_MAX_B64_LEN = 2_800_000
PREVIEW_IMAGE_MARGIN_CM = 0.45
DEFAULT_IMAGE_MARGIN_CM = 0.3


@dataclass(frozen=True)
class ImageLayout:
    max_width_cm: float
    max_height_cm: float = 24.0
    min_width_cm: float = 0.0
    min_width_source_threshold_cm: float = 0.0
    margin_cm: float = DEFAULT_IMAGE_MARGIN_CM


class ImageData:
    __slots__ = ("id", "png_bytes", "width_cm", "pixel_width", "pixel_height")

    def __init__(self, id: str, png_base64: str, width_cm: float = 15.5):
        self.id = id
        if len(png_base64) > IMAGE_MAX_B64_LEN:
            raise ValueError(f"Image {id}: base64 exceeds {IMAGE_MAX_B64_LEN} chars")
        self.png_bytes = base64.b64decode(png_base64)
        self.width_cm = width_cm

        img = Image.open(io.BytesIO(self.png_bytes))
        img.verify()
        w, h = img.size
        self.pixel_width = w
        self.pixel_height = h
        if w * h > IMAGE_MAX_PIXELS:
            raise ValueError(f"Image {id}: {w}x{h} = {w*h} pixels exceeds limit {IMAGE_MAX_PIXELS}")


def _iter_paragraphs(container):
    """递归遍历 document/cell 中的段落，覆盖表格单元格等容器。"""
    seen_cells: set[int] = set()
    for paragraph in getattr(container, "paragraphs", []):
        yield paragraph
    for table in getattr(container, "tables", []):
        for row in table.rows:
            for cell in row.cells:
                cell_key = id(cell._tc)
                if cell_key in seen_cells:
                    continue
                seen_cells.add(cell_key)
                yield from _iter_paragraphs(cell)


def resolve_image_width_cm(
    width_cm: float,
    style: str = "standard",
    layout: Optional[ImageLayout] = None,
    image_size: Optional[tuple[int, int]] = None,
) -> float:
    """根据导出样式解析图片插入宽度。"""
    if layout is not None:
        resolved = min(width_cm, layout.max_width_cm)
        should_enlarge = (
            layout.min_width_cm
            and width_cm >= layout.min_width_source_threshold_cm
            and resolved < layout.min_width_cm
        )
        if should_enlarge:
            resolved = min(layout.min_width_cm, layout.max_width_cm)
        if image_size:
            pixel_width, pixel_height = image_size
            if pixel_width > 0 and pixel_height > 0:
                ratio = pixel_height / pixel_width
                if ratio > 0:
                    resolved = min(resolved, layout.max_height_cm / ratio)
        resolved = max(1.0, resolved)
        return resolved

    if style == "preview":
        return min(width_cm, 19.0)
    return width_cm


def preprocess_markdown(md: str, images: list[dict]) -> tuple[str, Dict[str, ImageData]]:
    """预处理 markdown：验证图片数据，构建 id→ImageData 映射。

    将带 alt 的图表占位符规范化为 ![](id)，避免 Markdown 行内解析把它当成普通链接。
    返回 (原始 md, {id: ImageData})
    """
    md = PLACEHOLDER_IMAGE_MARKDOWN_PATTERN.sub(lambda m: f"\n\n![]({m.group(1)})\n\n", md)
    md = re.sub(r"\n{3,}", "\n\n", md).strip()
    image_map: Dict[str, ImageData] = {}
    for img in images:
        try:
            data = ImageData(
                id=img["id"],
                png_base64=img["pngBase64"],
                width_cm=img.get("widthCm", 15.5),
            )
            image_map[data.id] = data
        except Exception as e:
            logger.warning(f"[ImageInjector] Skipping image {img.get('id', '?')}: {e}")

    return md, image_map


def inject_images(
    doc_path: str,
    image_map: Dict[str, ImageData],
    style: str = "standard",
    layout: Optional[ImageLayout] = None,
) -> int:
    """在生成好的 DOCX 中，把占位符段落替换为图片。

    扫描所有段落，找到 ![](mdv__chart__xxx__) 格式的文本，替换为图片。
    返回成功注入的图片数量。
    """
    if not image_map:
        return 0

    doc = Document(doc_path)
    injected = 0

    for para in _iter_paragraphs(doc):
        text = para.text.strip()
        m = PLACEHOLDER_PATTERN.match(text)
        if not m:
            original_text = para.text
            inline_matches = list(PLACEHOLDER_IMAGE_MARKDOWN_PATTERN.finditer(original_text))
            if not inline_matches:
                continue

            for run in para.runs:
                run.clear()

            cursor = 0
            inserted_any = False
            for match in inline_matches:
                if match.start() > cursor:
                    para.add_run(original_text[cursor:match.start()])

                placeholder_id = match.group(1)
                img_data = image_map.get(placeholder_id)
                if img_data is None:
                    para.add_run(match.group(0))
                else:
                    run = para.add_run()
                    run.add_picture(
                        io.BytesIO(img_data.png_bytes),
                        width=Cm(resolve_image_width_cm(
                            img_data.width_cm,
                            style,
                            layout,
                            (img_data.pixel_width, img_data.pixel_height),
                        )),
                    )
                    injected += 1
                    inserted_any = True
                    logger.info(f"[ImageInjector] Injected {placeholder_id} ({len(img_data.png_bytes)} bytes)")
                cursor = match.end()

            if cursor < len(original_text):
                para.add_run(original_text[cursor:])

            if inserted_any:
                pf = para.paragraph_format
                pf.line_spacing = None
                pf.line_spacing_rule = None
                pf.first_line_indent = None
            continue

        placeholder_id = m.group(1)
        if placeholder_id not in image_map:
            logger.warning(f"[ImageInjector] Placeholder {placeholder_id} not found in image_map")
            continue

        img_data = image_map[placeholder_id]

        for run in para.runs:
            run.clear()

        run = para.add_run()
        run.add_picture(
            io.BytesIO(img_data.png_bytes),
            width=Cm(resolve_image_width_cm(
                img_data.width_cm,
                style,
                layout,
                (img_data.pixel_width, img_data.pixel_height),
            )),
        )
        injected += 1

        # 清除图片段落的固定行距和首行缩进（公文等样式的固定行距会压扁图片）
        pf = para.paragraph_format
        pf.line_spacing = None
        pf.line_spacing_rule = None
        pf.first_line_indent = None
        margin_cm = layout.margin_cm if layout is not None else (
            PREVIEW_IMAGE_MARGIN_CM if style == "preview" else DEFAULT_IMAGE_MARGIN_CM
        )
        pf.space_before = Cm(margin_cm)
        pf.space_after = Cm(margin_cm)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        logger.info(f"[ImageInjector] Injected {placeholder_id} ({len(img_data.png_bytes)} bytes)")

    doc.save(doc_path)
    return injected
