"""
Markdown → Word 文档生成器

支持多种导出格式：
- preview: 预览一致格式（接近 Markdown 预览 / PDF 导出）
- standard: 通用格式（清晰易读，适合日常使用）
- official: 正式公文格式（GB/T 9704-2012 严格版，四级标题体系）
- internal: 机关内部文件格式
- report: 调研/分析报告格式

Markdown 解析策略：使用正则逐行解析，避免引入额外依赖。
支持：标题(H1-H4)、段落、无序/有序列表、表格、代码块、粗体/斜体。
"""
import re
import logging
import subprocess
from typing import List, Optional, Dict
from dataclasses import dataclass, field
from enum import Enum

from app.presets import HeadingStyleDef, DOCX_PRESETS, VALID_STYLES, NON_PREVIEW_BLOCK_STYLES

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)

PREVIEW_FONT_CHAIN = (
    "PingFang SC",
    "Noto Sans CJK SC",
    "Microsoft YaHei",
    "SimHei",
    "Arial Unicode MS",
)

PREVIEW_BODY_LINE_SPACING = 1.45
PREVIEW_LIST_LINE_SPACING = 1.25
PREVIEW_TABLE_LINE_SPACING = 1.0
PREVIEW_TABLE_FONT_SIZE = Pt(8.0)
PREVIEW_TABLE_AFTER_SPACING = Pt(4)
PREVIEW_TABLE_CELL_MARGIN_TOP = 55
PREVIEW_TABLE_CELL_MARGIN_START = 110
PREVIEW_TABLE_CELL_MARGIN_BOTTOM = 55
PREVIEW_TABLE_CELL_MARGIN_END = 110

PREVIEW_MONO_FONT_CHAIN = (
    "Menlo",
    "Consolas",
    "Courier New",
    "Sarasa Mono SC",
    "Noto Sans Mono CJK SC",
)

PREVIEW_CONTENT_WIDTH_CM = 19.0


# ==================== 数据结构 ====================

class BlockType(Enum):
    HEADING = "heading"
    PARAGRAPH = "paragraph"
    UNORDERED_LIST = "unordered_list"
    ORDERED_LIST = "ordered_list"
    TABLE = "table"
    CODE_BLOCK = "code_block"
    HORIZONTAL_RULE = "horizontal_rule"
    BLOCKQUOTE = "blockquote"


@dataclass
class TextRun:
    """行内文本片段（支持粗体/斜体/代码/链接/引用角标）"""
    text: str
    bold: bool = False
    italic: bool = False
    code: bool = False
    link: str = ""               # 超链接 URL（非空时渲染为可点击链接）
    citation: bool = False       # 引用角标（上标样式）
    citation_ref_id: str = ""    # 关联的 ref_id（如 "ref_001"）


@dataclass
class Block:
    """文档块"""
    type: BlockType
    level: int = 0  # heading level (1-4)
    runs: List[TextRun] = field(default_factory=list)
    text: str = ""
    rows: List[List[str]] = field(default_factory=list)  # table rows


@dataclass
class MessageDTO:
    """消息数据传输对象"""
    role: str  # "user" | "assistant"
    content: str


# ==================== Markdown 解析 ====================

_RE_HEADING = re.compile(r"^(#{1,4})\s+(.+)$")
_RE_UL = re.compile(r"^(\s*)[-*+]\s+(.+)$")
_RE_OL = re.compile(r"^(\s*)\d+[.)]\s+(.+)$")
_RE_CODE_FENCE = re.compile(r"^```")
_RE_HR = re.compile(r"^(-{3,}|_{3,}|\*{3,})\s*$")
_RE_TABLE_ROW = re.compile(r"^\|(.+)\|$")
_RE_TABLE_SEP = re.compile(r"^\|[\s:|-]+\|$")
_RE_NOTE_PREFIX = re.compile(r"^\s*(?:\*\*)?\s*(注意|说明|注|Note|Warning)\s*[:：]\s*(?:\*\*)?\s*", re.IGNORECASE)
_RE_GFM_ALERT = re.compile(r"^\s*\[!(NOTE|WARNING|TIP|IMPORTANT)\]\s*", re.IGNORECASE)


def parse_inline(text: str, ref_id_to_index: Optional[Dict[str, int]] = None) -> List[TextRun]:
    """解析行内格式：**bold**, *italic*, `code`, [ref_XXX] 引用角标

    Args:
        ref_id_to_index: ref_id → 显示序号 映射（如 {"ref_001": 1}）。
                         提供时 [ref_XXX] 渲染为上标 [N]；未提供时保留原文。
    """
    runs: List[TextRun] = []
    # 引用角标 + 行内格式统一正则
    pattern = re.compile(
        r"(\[(?:ref:)?(ref_\d+)(?:[^\]]*)\])"  # [ref_001] / [ref:ref_001] / [ref_006评论区]
        r"|(\[(\d{1,3})\])(?!\()"               # [6] 纯数字（排除 markdown 链接）
        r"|(\[([^\]]+)\]\(([^)]+)\))"           # [text](url) markdown 链接
        r"|(?<!\w)(\*\*\*(.+?)\*\*\*)(?!\*)"         # ***bold italic***
        r"|(?<!\w)(\*\*(.+?)\*\*)(?!\*)"          # **bold**
        r"|(?<![*\w])(\*([^*]+?)\*)(?!\*)"         # *italic* (content must not contain *)
        r"|(`(.+?)`)"                           # `code`
    )
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            runs.append(TextRun(text=text[pos:m.start()]))

        if m.group(2):  # [ref_XXX] 格式
            ref_id = m.group(2)
            if ref_id_to_index and ref_id in ref_id_to_index:
                idx = ref_id_to_index[ref_id]
                runs.append(TextRun(text=f"[{idx}]", citation=True, citation_ref_id=ref_id))
            else:
                runs.append(TextRun(text=m.group(0)))
        elif m.group(4):  # [N] 纯数字
            num = int(m.group(4))
            ref_id = f"ref_{num:03d}"
            if ref_id_to_index and ref_id in ref_id_to_index:
                idx = ref_id_to_index[ref_id]
                runs.append(TextRun(text=f"[{idx}]", citation=True, citation_ref_id=ref_id))
            else:
                runs.append(TextRun(text=m.group(3)))
        elif m.group(5):  # [text](url) markdown 链接
            runs.append(TextRun(text=m.group(6), link=m.group(7)))
        elif m.group(9):  # bold italic
            runs.append(TextRun(text=m.group(9), bold=True, italic=True))
        elif m.group(11):  # bold
            runs.append(TextRun(text=m.group(11), bold=True))
        elif m.group(13):  # italic
            runs.append(TextRun(text=m.group(13), italic=True))
        elif m.group(15):  # code
            runs.append(TextRun(text=m.group(15), code=True))
        pos = m.end()

    if pos < len(text):
        runs.append(TextRun(text=text[pos:]))

    return runs or [TextRun(text=text)]


def parse_markdown(md_text: str, ref_id_to_index: Optional[Dict[str, int]] = None) -> List[Block]:
    """将 Markdown 文本解析为 Block 列表"""
    blocks: List[Block] = []
    lines = md_text.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]

        # 空行跳过
        if not line.strip():
            i += 1
            continue

        # 代码块
        if _RE_CODE_FENCE.match(line.strip()):
            code_lines = []
            i += 1
            while i < len(lines) and not _RE_CODE_FENCE.match(lines[i].strip()):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing fence
            blocks.append(Block(
                type=BlockType.CODE_BLOCK,
                text="\n".join(code_lines),
            ))
            continue

        # 水平线
        if _RE_HR.match(line.strip()):
            blocks.append(Block(type=BlockType.HORIZONTAL_RULE))
            i += 1
            continue

        # 标题
        m = _RE_HEADING.match(line.strip())
        if m:
            level = len(m.group(1))
            blocks.append(Block(
                type=BlockType.HEADING,
                level=level,
                runs=parse_inline(m.group(2).strip(), ref_id_to_index),
                text=m.group(2).strip(),
            ))
            i += 1
            continue

        # 表格
        if _RE_TABLE_ROW.match(line.strip()):
            rows = []
            while i < len(lines) and _RE_TABLE_ROW.match(lines[i].strip()):
                row_text = lines[i].strip()
                if _RE_TABLE_SEP.match(row_text):
                    i += 1
                    continue
                cells = [c.strip() for c in row_text.strip("|").split("|")]
                rows.append(cells)
                i += 1
            blocks.append(Block(type=BlockType.TABLE, rows=rows))
            continue

        # 无序列表
        m = _RE_UL.match(line)
        if m:
            blocks.append(Block(
                type=BlockType.UNORDERED_LIST,
                runs=parse_inline(m.group(2), ref_id_to_index),
                text=m.group(2),
            ))
            i += 1
            continue

        # 有序列表
        m = _RE_OL.match(line)
        if m:
            blocks.append(Block(
                type=BlockType.ORDERED_LIST,
                runs=parse_inline(m.group(2), ref_id_to_index),
                text=m.group(2),
            ))
            i += 1
            continue

        # Blockquote（> 开头的行，可能多行连续）
        if line.strip().startswith(">"):
            bq_lines = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                bq_lines.append(re.sub(r"^>\s?", "", lines[i].strip()))
                i += 1
            bq_text = "\n".join(l for l in bq_lines if l)
            blocks.append(Block(
                type=BlockType.BLOCKQUOTE,
                runs=parse_inline(bq_text, ref_id_to_index),
                text=bq_text,
            ))
            continue

        # 普通段落（可能多行连续）
        para_lines = [line.strip()]
        i += 1
        while i < len(lines) and lines[i].strip() and not any([
            _RE_HEADING.match(lines[i].strip()),
            _RE_UL.match(lines[i]),
            _RE_OL.match(lines[i]),
            _RE_CODE_FENCE.match(lines[i].strip()),
            _RE_HR.match(lines[i].strip()),
            _RE_TABLE_ROW.match(lines[i].strip()),
            lines[i].strip().startswith(">"),
        ]):
            para_lines.append(lines[i].strip())
            i += 1

        full_text = " ".join(para_lines)
        blocks.append(Block(
            type=BlockType.PARAGRAPH,
            runs=parse_inline(full_text, ref_id_to_index),
            text=full_text,
        ))

    return blocks


# ==================== Word 文档生成 ====================

def _add_hyperlink(
    paragraph,
    text: str,
    url: str,
    font_name: str,
    font_size: Pt,
    east_asia_font_name: Optional[str] = None,
    underline: bool = True,
):
    """向段落中添加可点击的超链接"""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = paragraph._element.makeelement(qn("w:hyperlink"), {qn("r:id"): r_id})
    new_run = paragraph._element.makeelement(qn("w:r"), {})
    rPr = paragraph._element.makeelement(qn("w:rPr"), {})
    rStyle = paragraph._element.makeelement(qn("w:rStyle"), {qn("w:val"): "Hyperlink"})
    rPr.append(rStyle)
    east_asia_font = east_asia_font_name or font_name
    rFonts = paragraph._element.makeelement(qn("w:rFonts"), {
        qn("w:ascii"): font_name, qn("w:hAnsi"): font_name, qn("w:eastAsia"): east_asia_font,
    })
    rPr.append(rFonts)
    sz = paragraph._element.makeelement(qn("w:sz"), {qn("w:val"): str(int(font_size.pt * 2))})
    rPr.append(sz)
    color = paragraph._element.makeelement(qn("w:color"), {qn("w:val"): "0563C1"})
    rPr.append(color)
    u = paragraph._element.makeelement(qn("w:u"), {qn("w:val"): "single" if underline else "none"})
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._element.append(hyperlink)


def _set_run_fonts(run, latin_font: str, east_asia_font: Optional[str] = None):
    run.font.name = latin_font
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), latin_font)
    r_fonts.set(qn("w:hAnsi"), latin_font)
    r_fonts.set(qn("w:eastAsia"), east_asia_font or latin_font)


def _apply_runs(
    paragraph,
    runs: List[TextRun],
    font_name: str = "宋体",
    font_size: Pt = Pt(12),
    east_asia_font_name: Optional[str] = None,
    code_font_name: str = "Courier New",
    code_east_asia_font_name: Optional[str] = None,
    link_underline: bool = True,
):
    """将 TextRun 列表写入 Word 段落"""
    for run_data in runs:
        if run_data.link:
            _add_hyperlink(
                paragraph,
                run_data.text,
                run_data.link,
                font_name,
                font_size,
                east_asia_font_name,
                underline=link_underline,
            )
            continue
        run = paragraph.add_run(run_data.text)
        _set_run_fonts(run, font_name, east_asia_font_name)
        run.font.size = font_size
        if run_data.bold:
            run.bold = True
        if run_data.italic:
            run.italic = True
        if run_data.code:
            _set_run_fonts(run, code_font_name, code_east_asia_font_name or code_font_name)
            run.font.size = Pt(font_size.pt - 1)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        if run_data.citation:
            run.font.size = Pt(max(font_size.pt - 3, 7))
            run.font.superscript = True
            run.font.color.rgb = RGBColor(0x20, 0x80, 0xF0)


def _set_paragraph_spacing(paragraph, before: Pt = Pt(0), after: Pt = Pt(0), line=None):
    """设置段落间距"""
    pf = paragraph.paragraph_format
    pf.space_before = before
    pf.space_after = after
    if line:
        pf.line_spacing = line


def _disable_paragraph_keep_constraints(paragraph):
    """避免 Word 内置 Heading 样式把标题和后续大图整组推到下一页。"""
    pf = paragraph.paragraph_format
    pf.keep_with_next = False
    pf.keep_together = False
    pf.page_break_before = False


def _available_font_families() -> set[str]:
    try:
        result = subprocess.run(
            ["fc-list", "--format=%{family}\n"],
            capture_output=True,
            text=True,
            timeout=5,
        )
        families: set[str] = set()
        for line in result.stdout.splitlines():
            for family in line.split(","):
                family = family.strip()
                if family:
                    families.add(family)
        return families
    except Exception:
        return set()


def _pick_font(candidates: tuple[str, ...], available: set[str]) -> str:
    if not available:
        return candidates[0]
    available_lower = {font.lower(): font for font in available}
    for candidate in candidates:
        if candidate.lower() in available_lower:
            return available_lower[candidate.lower()]
    return candidates[-1]


def _resolve_preview_fonts() -> tuple[str, str]:
    available = _available_font_families()
    return (
        _pick_font(PREVIEW_FONT_CHAIN, available),
        _pick_font(PREVIEW_MONO_FONT_CHAIN, available),
    )


def _set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _set_cell_margins(
    cell,
    top: int = PREVIEW_TABLE_CELL_MARGIN_TOP,
    start: int = PREVIEW_TABLE_CELL_MARGIN_START,
    bottom: int = PREVIEW_TABLE_CELL_MARGIN_BOTTOM,
    end: int = PREVIEW_TABLE_CELL_MARGIN_END,
):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def _set_table_borders(table, color: str = "D0D7DE", size: str = "4", val: str = "single"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), val)
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def _set_table_width(table, width_cm: float):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(int(width_cm * 567)))


def _preview_text_width_score(text: str) -> float:
    score = 0.0
    for ch in text:
        if ch.isspace():
            score += 0.4
        elif ord(ch) > 127:
            score += 1.8
        elif ch in "-_/.:()":
            score += 0.65
        else:
            score += 1.0
    return score


def _preview_plain_cell_text(text: str) -> str:
    text = re.sub(r"`([^`]*)`", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text.replace("**", "").replace("***", "").replace("*", "")


def _preview_column_widths_cm(rows: List[List[str]], total_width_cm: float = PREVIEW_CONTENT_WIDTH_CM) -> List[float]:
    if not rows:
        return []

    n_cols = max(len(row) for row in rows)
    if n_cols <= 0:
        return []

    weights: List[float] = []
    min_widths: List[float] = []
    for col_idx in range(n_cols):
        cells = [_preview_plain_cell_text(row[col_idx]) for row in rows if col_idx < len(row)]
        max_score = max((_preview_text_width_score(cell) for cell in cells), default=1.0)
        data_cells = cells[1:] if len(cells) > 1 else cells
        mostly_numeric = bool(data_cells) and all(re.fullmatch(r"[\d\s.,%A-Za-z/+-]+", cell or "") for cell in data_cells)
        if mostly_numeric:
            weights.append(3.0)
            min_widths.append(1.25)
        else:
            weights.append(max(5.0, min(max_score, 28.0)))
            min_widths.append(1.35)

    min_width_sum = sum(min_widths)
    if min_width_sum >= total_width_cm:
        return [total_width_cm / n_cols] * n_cols

    remaining_width = total_width_cm - min_width_sum
    total_weight = sum(weights) or 1.0
    widths = [min_width + remaining_width * (weight / total_weight) for min_width, weight in zip(min_widths, weights)]

    width_sum = sum(widths)
    if width_sum:
        scale = total_width_cm / width_sum
        widths = [width * scale for width in widths]
    return widths


def _preview_table_target_width_cm(rows: List[List[str]]) -> float:
    """让 preview 表格更接近浏览器 Markdown 的 auto-width，而不是所有表都铺满页面。"""
    if not rows:
        return PREVIEW_CONTENT_WIDTH_CM

    n_cols = max(len(row) for row in rows)
    if n_cols >= 4:
        return PREVIEW_CONTENT_WIDTH_CM

    estimated_widths: List[float] = []
    for col_idx in range(n_cols):
        cells = [_preview_plain_cell_text(row[col_idx]) for row in rows if col_idx < len(row)]
        max_score = max((_preview_text_width_score(cell) for cell in cells), default=1.0)
        estimated_widths.append(max(1.45, min(12.8, 1.0 + max_score * 0.12)))

    min_table_width = 5.0 if n_cols == 1 else 15.0
    target_width = max(min_table_width, sum(estimated_widths))
    return min(PREVIEW_CONTENT_WIDTH_CM, target_width)


def _adaptive_table_target_width_cm(rows: List[List[str]], content_width_cm: float) -> float:
    if not rows:
        return content_width_cm
    n_cols = max(len(row) for row in rows)
    if n_cols >= 4:
        return content_width_cm
    return min(content_width_cm, _preview_table_target_width_cm(rows))


def _add_table(
    doc,
    rows: List[List[str]],
    font_name: str = "宋体",
    font_size: Pt = Pt(10.5),
    ref_id_to_index: Optional[Dict[str, int]] = None,
    mode: str = "default",
    east_asia_font_name: Optional[str] = None,
    mono_font: str = "Courier New",
    mono_east_asia_font: Optional[str] = None,
):
    """添加 Word 表格（单元格内支持 **加粗**、*斜体*、`代码`、引用角标）"""
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols, style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    preview_table_width = PREVIEW_CONTENT_WIDTH_CM
    if mode == "preview":
        preview_table_width = _preview_table_target_width_cm(rows)
        table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        _set_table_width(table, preview_table_width)
        _set_table_borders(table)

    preview_cell_widths = _preview_column_widths_cm(rows, preview_table_width) if mode == "preview" and n_cols else []
    if preview_cell_widths:
        for j, width_cm in enumerate(preview_cell_widths):
            table.columns[j].width = Cm(width_cm)

    for i, row_data in enumerate(rows):
        if mode == "preview":
            _set_row_cant_split(table.rows[i])
        for j, cell_text in enumerate(row_data):
            if j < n_cols:
                cell = table.cell(i, j)
                if preview_cell_widths:
                    cell.width = Cm(preview_cell_widths[j])
                cell.text = ""
                p = cell.paragraphs[0]
                if mode == "preview":
                    _set_cell_margins(cell)
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing = PREVIEW_TABLE_LINE_SPACING
                    if i == 0:
                        _set_cell_shading(cell, "F6F8FA")
                # 表头行：整行加粗
                if i == 0:
                    run = p.add_run(cell_text)
                    _set_run_fonts(run, font_name, east_asia_font_name)
                    run.font.size = Pt(8.5) if mode == "preview" else font_size
                    run.bold = True
                else:
                    # 数据行：解析 inline markdown（加粗、斜体、引用等）
                    inline_runs = parse_inline(cell_text, ref_id_to_index)
                    _apply_runs(
                        p,
                        inline_runs,
                        font_name=font_name,
                        font_size=font_size,
                        east_asia_font_name=east_asia_font_name,
                        code_font_name=mono_font,
                        code_east_asia_font_name=mono_east_asia_font,
                        link_underline=mode != "preview",
                    )


def _add_styled_table(
    doc,
    rows: List[List[str]],
    table_style,
    font_name: str,
    ref_id_to_index: Optional[Dict[str, int]] = None,
    east_asia_font_name: Optional[str] = None,
    mono_font: str = "Courier New",
    mono_east_asia_font: Optional[str] = None,
):
    """添加非 preview 样式表格。preview 继续走 _add_table(mode="preview")。"""
    if not rows:
        return

    n_cols = max(len(row) for row in rows)
    table_width = table_style.content_width_cm
    if table_style.adaptive_width:
        table_width = _adaptive_table_target_width_cm(rows, table_width)

    table = doc.add_table(rows=len(rows), cols=n_cols, style="Table Grid")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT if table_style.alignment == "left" else WD_TABLE_ALIGNMENT.CENTER
    _set_table_width(table, table_width)
    _set_table_borders(table, color=table_style.border_color, size=table_style.border_size)

    column_widths = _preview_column_widths_cm(rows, table_width) if n_cols else []
    for j, width_cm in enumerate(column_widths):
        table.columns[j].width = Cm(width_cm)

    for i, row_data in enumerate(rows):
        _set_row_cant_split(table.rows[i])
        for j, cell_text in enumerate(row_data):
            if j >= n_cols:
                continue
            cell = table.cell(i, j)
            if j < len(column_widths):
                cell.width = Cm(column_widths[j])
            cell.text = ""
            _set_cell_margins(
                cell,
                top=table_style.cell_margin_top,
                start=table_style.cell_margin_start,
                bottom=table_style.cell_margin_bottom,
                end=table_style.cell_margin_end,
            )
            if i == 0 and table_style.header_fill:
                _set_cell_shading(cell, table_style.header_fill)

            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = table_style.line_spacing
            font_size = Pt(table_style.header_font_size if i == 0 else table_style.body_font_size)

            if i == 0:
                run = paragraph.add_run(cell_text)
                _set_run_fonts(run, font_name, east_asia_font_name)
                run.font.size = font_size
                run.bold = True
            else:
                _apply_runs(
                    paragraph,
                    parse_inline(cell_text, ref_id_to_index),
                    font_name=font_name,
                    font_size=font_size,
                    east_asia_font_name=east_asia_font_name,
                    code_font_name=mono_font,
                    code_east_asia_font_name=mono_east_asia_font,
                )


def _add_preview_table_gap(doc):
    """模拟浏览器 Markdown 表格的 margin-bottom，避免表格紧贴后续正文。"""
    para = doc.add_paragraph()
    _set_paragraph_spacing(para, before=Pt(0), after=PREVIEW_TABLE_AFTER_SPACING, line=Pt(1))
    run = para.add_run("")
    run.font.size = Pt(1)


def _add_table_gap(doc, after: Pt = Pt(4)):
    para = doc.add_paragraph()
    _set_paragraph_spacing(para, before=Pt(0), after=after, line=Pt(1))
    run = para.add_run("")
    run.font.size = Pt(1)


def _extract_note_text(text: str) -> Optional[str]:
    stripped = text.strip()
    alert = _RE_GFM_ALERT.match(stripped)
    if alert:
        return stripped[alert.end():].strip()
    prefix = _RE_NOTE_PREFIX.match(stripped)
    if prefix:
        return stripped[prefix.end():].strip()
    return None


def _add_box_callout(
    doc,
    text: str,
    font_name: str,
    body_size: Pt,
    fill: str,
    width_cm: float,
    ref_id_to_index: Optional[Dict[str, int]] = None,
    east_asia_font_name: Optional[str] = None,
    mono_font: str = "Courier New",
    mono_east_asia_font: Optional[str] = None,
    link_underline: bool = True,
):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_table_width(table, width_cm)
    _set_table_borders(table, color=fill, size="0", val="nil")

    cell = table.cell(0, 0)
    cell.width = Cm(width_cm)
    _set_cell_shading(cell, fill)
    _set_cell_margins(cell, top=80, start=130, bottom=80, end=130)

    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = PREVIEW_LIST_LINE_SPACING

    for li, line in enumerate(text.split("\n")):
        _apply_runs(
            para,
            parse_inline(line, ref_id_to_index),
            font_name=font_name,
            font_size=Pt(max(body_size.pt - 0.5, 8.0)),
            east_asia_font_name=east_asia_font_name,
            code_font_name=mono_font,
            code_east_asia_font_name=mono_east_asia_font,
            link_underline=link_underline,
        )
        if li < len(text.split("\n")) - 1:
            para.add_run().add_break()

    for run in para.runs:
        run.font.color.rgb = RGBColor(0x57, 0x60, 0x6A)

    _add_table_gap(doc, after=Pt(4))


def _add_non_preview_callout(
    doc,
    text: str,
    font_name: str,
    body_size: Pt,
    block_style,
    ref_id_to_index: Optional[Dict[str, int]] = None,
    east_asia_font_name: Optional[str] = None,
    mono_font: str = "Courier New",
    mono_east_asia_font: Optional[str] = None,
    link_underline: bool = True,
):
    note_text = _extract_note_text(text)
    if block_style.callout.mode == "official":
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.left_indent = Cm(0.74)
        display_text = f"{block_style.callout.note_prefix}{note_text}" if note_text is not None else text
        _apply_runs(
            para,
            parse_inline(display_text, ref_id_to_index),
            font_name=font_name,
            font_size=body_size,
            east_asia_font_name=east_asia_font_name,
            code_font_name=mono_font,
            code_east_asia_font_name=mono_east_asia_font,
            link_underline=link_underline,
        )
        _set_paragraph_spacing(para, before=Pt(2), after=Pt(2))
        return

    if note_text is not None:
        _add_box_callout(
            doc,
            text,
            font_name,
            body_size,
            fill=block_style.callout.fill or "F6F8FA",
            width_cm=block_style.table.content_width_cm,
            ref_id_to_index=ref_id_to_index,
            east_asia_font_name=east_asia_font_name,
            mono_font=mono_font,
            mono_east_asia_font=mono_east_asia_font,
            link_underline=link_underline,
        )
        return

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.left_indent = Cm(0.74)
    for li, line in enumerate(text.split("\n")):
        _apply_runs(
            para,
            parse_inline(line, ref_id_to_index),
            font_name=font_name,
            font_size=Pt(max(body_size.pt + block_style.callout.font_size_delta, 8.0)),
            east_asia_font_name=east_asia_font_name,
            code_font_name=mono_font,
            code_east_asia_font_name=mono_east_asia_font,
            link_underline=link_underline,
        )
        if li < len(text.split("\n")) - 1:
            para.add_run().add_break()
    _set_paragraph_spacing(para, before=Pt(2), after=Pt(2))


def _add_preview_callout(
    doc,
    text: str,
    font_name: str,
    body_size: Pt,
    ref_id_to_index: Optional[Dict[str, int]] = None,
    east_asia_font_name: Optional[str] = None,
    mono_font: str = "Courier New",
    mono_east_asia_font: Optional[str] = None,
    link_underline: bool = False,
):
    """用一列表格模拟 Markdown 预览里的浅灰提示块，避免 Word blockquote 竖线样式。"""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_table_width(table, PREVIEW_CONTENT_WIDTH_CM)
    _set_table_borders(table, color="F6F8FA", size="0", val="nil")

    cell = table.cell(0, 0)
    cell.width = Cm(PREVIEW_CONTENT_WIDTH_CM)
    _set_cell_shading(cell, "F6F8FA")
    _set_cell_margins(cell, top=80, start=130, bottom=80, end=130)

    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = PREVIEW_LIST_LINE_SPACING

    lines = text.split("\n")
    callout_font_size = Pt(max(body_size.pt - 0.5, 8.0))
    for li, line in enumerate(lines):
        _apply_runs(
            para,
            parse_inline(line, ref_id_to_index),
            font_name=font_name,
            font_size=callout_font_size,
            east_asia_font_name=east_asia_font_name,
            code_font_name=mono_font,
            code_east_asia_font_name=mono_east_asia_font,
            link_underline=link_underline,
        )
        if li < len(lines) - 1:
            para.add_run().add_break()

    for run in para.runs:
        run.font.color.rgb = RGBColor(0x57, 0x60, 0x6A)

    _add_preview_table_gap(doc)


def generate_standard_docx(
    title: str,
    messages: List[MessageDTO],
    output_path: str,
    footer_text: str = "由终身教育智能体生成",
):
    """
    通用格式 Word 导出

    样式：宋体正文、标题加粗、适中行距、清晰的角色分隔
    """
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # 文档标题
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    title_run.font.name = "黑体"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    title_run.font.size = Pt(18)
    title_run.bold = True
    _set_paragraph_spacing(title_para, before=Pt(12), after=Pt(12))

    # 分隔线
    doc.add_paragraph("").runs  # 空行

    # 渲染每条消息
    for msg in messages:
        # 角色标签
        role_name = "用户" if msg.role == "user" else "助手"
        role_para = doc.add_paragraph()
        role_run = role_para.add_run(f"【{role_name}】")
        role_run.font.name = "黑体"
        role_run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        role_run.font.size = Pt(12)
        role_run.bold = True
        if msg.role == "user":
            role_run.font.color.rgb = RGBColor(0x20, 0x80, 0xF0)
        else:
            role_run.font.color.rgb = RGBColor(0x18, 0xA0, 0x58)
        _set_paragraph_spacing(role_para, before=Pt(12), after=Pt(4))

        # 解析并渲染消息内容
        blocks = parse_markdown(msg.content)
        _render_blocks(doc, blocks, font_name="宋体", body_size=Pt(12))

    # 页脚
    doc.add_paragraph("")
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(footer_text)
    footer_run.font.name = "宋体"
    footer_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.save(output_path)
    logger.info(f"[DocxExport] 标准格式导出完成: {output_path}")


def generate_official_docx(
    title: str,
    messages: List[MessageDTO],
    output_path: str,
    footer_text: str = "由终身教育智能体生成",
):
    """
    行政公文标准格式 Word 导出（GB/T 9704-2012 近似）

    排版规范：
    - 标题：黑体 二号（22pt），居中
    - 正文：宋体 三号（16pt），两端对齐
    - 行距：28磅固定行距
    - 页边距：上3.7cm 下3.5cm 左2.8cm 右2.6cm
    """
    doc = Document()

    # 页面设置（公文标准页边距）
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)

    # 文档标题：黑体 二号
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    title_run.font.name = "黑体"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    title_run.font.size = Pt(22)
    title_run.bold = True
    _set_paragraph_spacing(title_para, before=Pt(0), after=Pt(16), line=Pt(28))

    # 空行分隔
    sep = doc.add_paragraph()
    _set_paragraph_spacing(sep, line=Pt(28))

    # 公文格式：只渲染 assistant 消息的内容作为正文
    # 用户消息不显示（公文是最终输出文件）
    assistant_msgs = [m for m in messages if m.role == "assistant"]
    if not assistant_msgs:
        assistant_msgs = messages  # fallback: 没有 assistant 消息时显示所有

    for msg in assistant_msgs:
        blocks = parse_markdown(msg.content)
        _render_blocks(
            doc, blocks,
            font_name="宋体",
            body_size=Pt(16),
            heading_font="黑体",
            line_spacing=Pt(28),
            first_line_indent=Cm(0.74),  # 两字符缩进
            align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        )

    # 页脚
    doc.add_paragraph("")
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(footer_text)
    footer_run.font.name = "宋体"
    footer_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.save(output_path)
    logger.info(f"[DocxExport] 公文格式导出完成: {output_path}")


def _render_blocks(
    doc: Document,
    blocks: List[Block],
    font_name: str = "宋体",
    body_size: Pt = Pt(12),
    heading_font: str = "黑体",
    line_spacing=None,
    first_line_indent: Optional[Cm] = None,
    align: int = WD_ALIGN_PARAGRAPH.LEFT,
    heading_styles: Optional[Dict[int, "HeadingStyleDef"]] = None,
    ref_id_to_index: Optional[Dict[str, int]] = None,
    table_mode: str = "default",
    code_mode: str = "default",
    mono_font: str = "Courier New",
    east_asia_font_name: Optional[str] = None,
    mono_east_asia_font: Optional[str] = None,
    heading_spacing_before: Pt = Pt(12),
    heading_spacing_after: Pt = Pt(6),
    list_left_indent: Cm = Cm(1.2),
    list_hanging_indent: Cm = Cm(0.5),
    list_spacing_before: Pt = Pt(1),
    list_spacing_after: Pt = Pt(1),
    list_line_spacing=None,
    link_underline: bool = True,
    block_style=None,
):
    """将解析后的 Block 列表渲染到 Word 文档

    Args:
        heading_styles: 按级别配置标题样式 {level: HeadingStyleDef}。
                        如果提供，则覆盖 heading_font 和 heading_sizes 的线性递减逻辑。
    """

    heading_sizes = {
        1: Pt(body_size.pt + 6),
        2: Pt(body_size.pt + 4),
        3: Pt(body_size.pt + 2),
        4: Pt(body_size.pt + 1),
    }

    list_counter = 0  # 有序列表计数器
    prev_was_ol = False
    preview_mode = table_mode == "preview" or code_mode == "preview"

    for idx, block in enumerate(blocks):
        if block.type == BlockType.HEADING:
            heading_level = min(block.level, 4)
            if heading_styles and block.level in heading_styles:
                hs = heading_styles[block.level]
                h_font = hs.font if isinstance(hs, HeadingStyleDef) else hs.get("font", heading_font)
                h_size = Pt(hs.size if isinstance(hs, HeadingStyleDef) else hs.get("size", body_size.pt))
                h_bold = hs.bold if isinstance(hs, HeadingStyleDef) else hs.get("bold", False)
                para = doc.add_heading(level=heading_level)
                para.clear()
                _apply_runs(
                    para,
                    block.runs,
                    font_name=h_font,
                    font_size=h_size,
                    east_asia_font_name=east_asia_font_name,
                    code_font_name=mono_font,
                    code_east_asia_font_name=mono_east_asia_font,
                    link_underline=link_underline,
                )
                for run in para.runs:
                    run.bold = h_bold
                    run.italic = False
            else:
                size = heading_sizes.get(block.level, body_size)
                para = doc.add_heading(level=heading_level)
                para.clear()
                _apply_runs(
                    para,
                    block.runs,
                    font_name=heading_font,
                    font_size=size,
                    east_asia_font_name=east_asia_font_name,
                    code_font_name=mono_font,
                    code_east_asia_font_name=mono_east_asia_font,
                    link_underline=link_underline,
                )
                for run in para.runs:
                    run.bold = True
                    run.italic = False

            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in para.runs:
                run.italic = False
                run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
                # 清除 Word 内置 Heading 的主题色（否则主题色覆盖 RGB）
                rPr = run._element.get_or_add_rPr()
                color_el = rPr.find(qn('w:color'))
                if color_el is not None:
                    for attr in ['themeColor', 'themeShade', 'themeTint']:
                        color_el.attrib.pop(qn(f'w:{attr}'), None)
            _set_paragraph_spacing(para, before=heading_spacing_before, after=heading_spacing_after, line=line_spacing)
            _disable_paragraph_keep_constraints(para)

        elif block.type == BlockType.PARAGRAPH:
            para = doc.add_paragraph()
            para.alignment = align
            _apply_runs(
                para,
                block.runs,
                font_name=font_name,
                font_size=body_size,
                east_asia_font_name=east_asia_font_name,
                code_font_name=mono_font,
                code_east_asia_font_name=mono_east_asia_font,
                link_underline=link_underline,
            )
            _set_paragraph_spacing(para, before=Pt(2), after=Pt(2), line=line_spacing)
            if first_line_indent:
                para.paragraph_format.first_line_indent = first_line_indent

        elif block.type == BlockType.UNORDERED_LIST:
            para = doc.add_paragraph()
            para.alignment = align
            para.paragraph_format.left_indent = list_left_indent
            para.paragraph_format.first_line_indent = -list_hanging_indent
            bullet_run = para.add_run("\u2022  ")
            _set_run_fonts(bullet_run, font_name, east_asia_font_name)
            bullet_run.font.size = body_size
            _apply_runs(
                para,
                block.runs,
                font_name=font_name,
                font_size=body_size,
                east_asia_font_name=east_asia_font_name,
                code_font_name=mono_font,
                code_east_asia_font_name=mono_east_asia_font,
                link_underline=link_underline,
            )
            _set_paragraph_spacing(para, before=list_spacing_before, after=list_spacing_after, line=list_line_spacing or line_spacing)
            prev_was_ol = False

        elif block.type == BlockType.ORDERED_LIST:
            if not prev_was_ol:
                list_counter = 0
            list_counter += 1
            para = doc.add_paragraph()
            para.alignment = align
            para.paragraph_format.left_indent = list_left_indent
            para.paragraph_format.first_line_indent = -list_hanging_indent
            num_run = para.add_run(f"{list_counter}.  ")
            _set_run_fonts(num_run, font_name, east_asia_font_name)
            num_run.font.size = body_size
            _apply_runs(
                para,
                block.runs,
                font_name=font_name,
                font_size=body_size,
                east_asia_font_name=east_asia_font_name,
                code_font_name=mono_font,
                code_east_asia_font_name=mono_east_asia_font,
                link_underline=link_underline,
            )
            _set_paragraph_spacing(para, before=list_spacing_before, after=list_spacing_after, line=list_line_spacing or line_spacing)
            prev_was_ol = True

        elif block.type == BlockType.TABLE:
            if block_style is not None:
                _add_styled_table(
                    doc,
                    block.rows,
                    table_style=block_style.table,
                    font_name=font_name,
                    ref_id_to_index=ref_id_to_index,
                    east_asia_font_name=east_asia_font_name,
                    mono_font=mono_font,
                    mono_east_asia_font=mono_east_asia_font,
                )
                _add_table_gap(doc, after=Pt(block_style.table.gap_after_pt))
                continue

            table_font_size = PREVIEW_TABLE_FONT_SIZE if table_mode == "preview" else Pt(body_size.pt - 1.5)
            _add_table(
                doc,
                block.rows,
                font_name=font_name,
                font_size=table_font_size,
                ref_id_to_index=ref_id_to_index,
                mode=table_mode,
                east_asia_font_name=east_asia_font_name,
                mono_font=mono_font,
                mono_east_asia_font=mono_east_asia_font,
            )
            if table_mode == "preview":
                _add_preview_table_gap(doc)

        elif block.type == BlockType.CODE_BLOCK:
            code_font = mono_font if code_mode == "preview" else "Courier New"
            code_font_size = Pt(
                block_style.code.font_size
                if block_style is not None
                else (8.5 if code_mode == "preview" else 9)
            )
            code_fill = (
                block_style.code.fill
                if block_style is not None
                else ("F6F8FA" if code_mode == "preview" else "F5F5F5")
            )
            code_line_spacing = block_style.code.line_spacing if block_style is not None else 1.0
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = para.add_run(block.text)
            _set_run_fonts(run, code_font, mono_east_asia_font if code_mode == "preview" else code_font)
            run.font.size = code_font_size
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            shading = para._element.makeelement(qn("w:shd"), {
                qn("w:val"): "clear",
                qn("w:color"): "auto",
                qn("w:fill"): code_fill,
            })
            para._element.get_or_add_pPr().append(shading)
            _set_paragraph_spacing(para, before=Pt(4), after=Pt(4))
            para.paragraph_format.line_spacing = code_line_spacing

        elif block.type == BlockType.BLOCKQUOTE:
            if block_style is not None:
                _add_non_preview_callout(
                    doc,
                    block.text,
                    font_name=font_name,
                    body_size=body_size,
                    block_style=block_style,
                    ref_id_to_index=ref_id_to_index,
                    east_asia_font_name=east_asia_font_name,
                    mono_font=mono_font,
                    mono_east_asia_font=mono_east_asia_font,
                    link_underline=link_underline,
                )
                continue

            if preview_mode:
                _add_preview_callout(
                    doc,
                    block.text,
                    font_name=font_name,
                    body_size=body_size,
                    ref_id_to_index=ref_id_to_index,
                    east_asia_font_name=east_asia_font_name,
                    mono_font=mono_font,
                    mono_east_asia_font=mono_east_asia_font,
                    link_underline=link_underline,
                )
                continue

            para = doc.add_paragraph()
            para.alignment = align
            bq_lines = block.text.split("\n")
            for li, bq_line in enumerate(bq_lines):
                line_runs = parse_inline(bq_line, ref_id_to_index)
                _apply_runs(
                    para,
                    line_runs,
                    font_name=font_name,
                    font_size=body_size,
                    east_asia_font_name=east_asia_font_name,
                    code_font_name=mono_font,
                    code_east_asia_font_name=mono_east_asia_font,
                    link_underline=link_underline,
                )
                if li < len(bq_lines) - 1:
                    br_run = para.add_run()
                    br_run.add_break()
            for run in para.runs:
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            para.paragraph_format.left_indent = Cm(1)
            bq_shd = para._element.makeelement(qn("w:shd"), {
                qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): "F6F8FA",
            })
            para._element.get_or_add_pPr().append(bq_shd)
            bq_border = para._element.makeelement(qn("w:pBdr"), {})
            left_border = para._element.makeelement(qn("w:left"), {
                qn("w:val"): "single", qn("w:sz"): "24",
                qn("w:space"): "4", qn("w:color"): "CCCCCC",
            })
            bq_border.append(left_border)
            para._element.get_or_add_pPr().append(bq_border)
            _set_paragraph_spacing(para, before=Pt(4), after=Pt(4), line=line_spacing)

        elif block.type == BlockType.HORIZONTAL_RULE:
            next_block = blocks[idx + 1] if idx + 1 < len(blocks) else None
            if not preview_mode and next_block and next_block.type == BlockType.HEADING:
                continue
            para = doc.add_paragraph()
            if preview_mode:
                p_bdr = OxmlElement("w:pBdr")
                bottom = OxmlElement("w:bottom")
                bottom.set(qn("w:val"), "single")
                bottom.set(qn("w:sz"), "4")
                bottom.set(qn("w:space"), "1")
                bottom.set(qn("w:color"), "EAECEF")
                p_bdr.append(bottom)
                para._element.get_or_add_pPr().append(p_bdr)
                _set_paragraph_spacing(para, before=Pt(8), after=Pt(8), line=Pt(1))
            else:
                _set_paragraph_spacing(para, before=Pt(4), after=Pt(4))

        # 重置有序列表计数器
        if block.type != BlockType.ORDERED_LIST:
            prev_was_ol = False


# ==================== 预设体系（已提取到 app/presets.py）====================


def _build_ref_id_to_index(references: Optional[List[dict]]) -> Optional[Dict[str, int]]:
    """从 references 列表构建 ref_id → 显示序号 映射"""
    if not references:
        return None
    mapping: Dict[str, int] = {}
    for i, ref in enumerate(references, 1):
        ref_id = ref.get('id', '')
        if ref_id:
            mapping[ref_id] = i
    return mapping if mapping else None


def _add_reference_list(
    doc: Document,
    references: List[dict],
    font_name: str = "宋体",
    body_size: Pt = Pt(12),
    heading_font: str = "黑体",
    line_spacing: Optional[Pt] = None,
):
    """在文档末尾添加"参考来源"章节"""
    if not references:
        return

    # 分隔线
    doc.add_paragraph("")

    # 标题
    heading_para = doc.add_paragraph()
    heading_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading_run = heading_para.add_run("参考来源")
    heading_run.font.name = heading_font
    heading_run._element.rPr.rFonts.set(qn("w:eastAsia"), heading_font)
    heading_run.font.size = Pt(body_size.pt + 2)
    heading_run.bold = True
    _set_paragraph_spacing(heading_para, before=Pt(12), after=Pt(6), line=line_spacing)

    # 逐条列出
    for i, ref in enumerate(references, 1):
        title = ref.get('title', ref.get('filename', f'来源 {i}'))
        url = ref.get('url', '')
        snippet = ref.get('snippet', '')

        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # 序号
        idx_run = para.add_run(f"[{i}] ")
        idx_run.font.name = font_name
        idx_run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        idx_run.font.size = Pt(body_size.pt - 1)
        idx_run.bold = True

        # 标题
        title_run = para.add_run(title)
        title_run.font.name = font_name
        title_run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        title_run.font.size = Pt(body_size.pt - 1)

        # URL
        if url:
            url_run = para.add_run(f"  {url}")
            url_run.font.name = font_name
            url_run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
            url_run.font.size = Pt(body_size.pt - 2)
            url_run.font.color.rgb = RGBColor(0x20, 0x80, 0xF0)

        _set_paragraph_spacing(para, before=Pt(1), after=Pt(1), line=line_spacing)


def generate_docx_from_content(
    content: str,
    output_path: str,
    *,
    style: str = "standard",
    title: str = None,
    footer_text: str = "由终身教育智能体生成",
    references: Optional[List[dict]] = None,
    reference_docx_path: Optional[str] = None,
):
    """
    从 Markdown 内容生成格式化 Word 文档（统一入口）

    Args:
        content: Markdown 格式的文档内容
        output_path: 输出文件路径
        style: 排版预设 ID（preview/official/internal/report/standard）
        title: 文档标题，为空时自动从内容首个 # 标题提取
        footer_text: 页脚文字
        references: 引用列表（来自 ArtifactVersion.metadata_），用于渲染引用角标和参考来源
    """
    # 白名单校验
    if style not in VALID_STYLES:
        logger.warning(f"[DocxExport] 无效的 style 参数: {style!r}，回退到 standard")
        style = "standard"

    # 构建引用映射
    ref_id_to_index = _build_ref_id_to_index(references)

    if style == "preview":
        _generate_preview_from_content(
            content,
            output_path,
            title,
            footer_text,
            references=references,
            ref_id_to_index=ref_id_to_index,
            reference_docx_path=reference_docx_path,
        )
        return

    # standard 模式：基础渲染
    if style == "standard":
        _generate_standard_from_content(
            content,
            output_path,
            title,
            footer_text,
            references=references,
            ref_id_to_index=ref_id_to_index,
            reference_docx_path=reference_docx_path,
        )
        return

    preset = DOCX_PRESETS[style]

    # 提取标题
    if not title:
        for line in content.split('\n'):
            stripped = line.strip()
            if stripped.startswith('# '):
                title = stripped.lstrip('#').strip()
                break
        if not title:
            title = preset.get("display_name", "文档")

    doc = Document(reference_docx_path) if reference_docx_path else Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    margins = preset["page_margins"]
    section.top_margin = Cm(margins["top"])
    section.bottom_margin = Cm(margins["bottom"])
    section.left_margin = Cm(margins["left"])
    section.right_margin = Cm(margins["right"])

    # 文档标题
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    title_font_name = preset["title_font"]
    title_run.font.name = title_font_name
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), title_font_name)
    title_run.font.size = Pt(preset["title_size"])
    title_run.bold = True

    # 标题行距
    line_sp = Pt(preset["line_spacing"]) if "line_spacing" in preset else None
    title_after = Pt(10 if style == "official" else 8)
    _set_paragraph_spacing(title_para, before=Pt(0), after=title_after, line=line_sp)

    # 正文渲染参数
    body_font = preset["body_font"]
    body_size = Pt(preset["body_size"])
    first_indent = Cm(preset["first_line_indent"]) if preset.get("first_line_indent") else None
    align_map = {"justify": WD_ALIGN_PARAGRAPH.JUSTIFY, "left": WD_ALIGN_PARAGRAPH.LEFT}
    align_val = align_map.get(preset.get("align", "justify"), WD_ALIGN_PARAGRAPH.JUSTIFY)

    # 行距处理：固定行距 vs 倍数行距
    if "line_spacing" in preset:
        line_spacing_val = Pt(preset["line_spacing"])
    else:
        line_spacing_val = None

    # 解析并渲染（去掉第一个 H1 标题，因为已单独渲染为文档标题）
    blocks = parse_markdown(content, ref_id_to_index)
    if blocks and blocks[0].type == BlockType.HEADING and blocks[0].level == 1:
        blocks = blocks[1:]
    _render_blocks(
        doc, blocks,
        font_name=body_font,
        body_size=body_size,
        heading_font=preset.get("title_font", "黑体"),
        line_spacing=line_spacing_val,
        first_line_indent=first_indent,
        align=align_val,
        heading_styles=preset.get("heading_styles"),
        ref_id_to_index=ref_id_to_index,
        heading_spacing_before=Pt(8 if style == "official" else 6),
        heading_spacing_after=Pt(4),
        list_spacing_before=Pt(0),
        list_spacing_after=Pt(0),
        block_style=NON_PREVIEW_BLOCK_STYLES.get(style),
    )

    # 倍数行距需要额外处理（_render_blocks 只支持固定行距参数）
    if "line_spacing_multiple" in preset:
        from docx.shared import Emu
        for para in doc.paragraphs[1:]:  # 跳过标题
            pf = para.paragraph_format
            if pf.line_spacing is None:
                pf.line_spacing = preset["line_spacing_multiple"]

    # 参考来源
    if references:
        _add_reference_list(
            doc, references,
            font_name=body_font, body_size=body_size,
            heading_font=preset.get("title_font", "黑体"),
            line_spacing=line_spacing_val,
        )

    # 页脚
    doc.add_paragraph("")
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(footer_text)
    footer_run.font.name = "宋体"
    footer_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.save(output_path)
    logger.info(f"[DocxExport] {style} 格式导出完成: {output_path}")


def _generate_preview_from_content(
    content: str,
    output_path: str,
    title: str = None,
    footer_text: str = "由 MD Viewer 生成",
    references: Optional[List[dict]] = None,
    ref_id_to_index: Optional[Dict[str, int]] = None,
    reference_docx_path: Optional[str] = None,
):
    """预览一致格式渲染：A4、窄边距、紧凑表格与接近预览的正文密度。"""
    if not title:
        for line in content.split("\n"):
            stripped = line.strip()
            if stripped.startswith("# "):
                title = stripped.lstrip("#").strip()
                break

    doc = Document(reference_docx_path) if reference_docx_path else Document()

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)

    body_font, mono_font = _resolve_preview_fonts()
    body_size = Pt(10)

    if title:
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        title_run = title_para.add_run(title)
        _set_run_fonts(title_run, body_font)
        title_run.font.size = Pt(18)
        title_run.bold = False
        title_run.italic = False
        title_run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
        _set_paragraph_spacing(title_para, before=Pt(0), after=Pt(10))

    blocks = parse_markdown(content, ref_id_to_index)
    if blocks and blocks[0].type == BlockType.HEADING and blocks[0].level == 1:
        blocks = blocks[1:]

    _render_blocks(
        doc,
        blocks,
        font_name=body_font,
        body_size=body_size,
        heading_font=body_font,
        line_spacing=PREVIEW_BODY_LINE_SPACING,
        first_line_indent=None,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        heading_styles={
            1: HeadingStyleDef(body_font, 18, False),
            2: HeadingStyleDef(body_font, 14, False),
            3: HeadingStyleDef(body_font, 12, True),
            4: HeadingStyleDef(body_font, 11, True),
        },
        ref_id_to_index=ref_id_to_index,
        table_mode="preview",
        code_mode="preview",
        mono_font=mono_font,
        heading_spacing_before=Pt(8),
        heading_spacing_after=Pt(4),
        list_left_indent=Cm(0.85),
        list_hanging_indent=Cm(0.32),
        list_spacing_before=Pt(0),
        list_spacing_after=Pt(0),
        list_line_spacing=PREVIEW_LIST_LINE_SPACING,
        link_underline=False,
    )

    if references:
        _add_reference_list(
            doc,
            references,
            font_name=body_font,
            body_size=body_size,
            heading_font=body_font,
            line_spacing=PREVIEW_BODY_LINE_SPACING,
        )

    if footer_text:
        doc.add_paragraph("")
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run(footer_text)
        _set_run_fonts(footer_run, body_font)
        footer_run.font.size = Pt(8)
        footer_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.save(output_path)
    logger.info(f"[DocxExport] preview 格式导出完成: {output_path}")


def _generate_standard_from_content(
    content: str,
    output_path: str,
    title: str = None,
    footer_text: str = "由终身教育智能体生成",
    references: Optional[List[dict]] = None,
    ref_id_to_index: Optional[Dict[str, int]] = None,
    reference_docx_path: Optional[str] = None,
):
    """标准格式渲染（基础 markdown→word，与原 FileService._write_docx 等效）"""
    if not title:
        for line in content.split('\n'):
            stripped = line.strip()
            if stripped.startswith('# '):
                title = stripped.lstrip('#').strip()
                break

    doc = Document(reference_docx_path) if reference_docx_path else Document()

    # 标准页面设置
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.3)
    section.bottom_margin = Cm(2.3)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # standard 模式用现代无衬线字体（与 PDF 导出风格一致）
    std_body_font = "微软雅黑"
    std_heading_font = "微软雅黑"
    std_body_size = Pt(11)

    # 标题（如果有）
    if title:
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(title)
        title_run.font.name = std_heading_font
        title_run._element.rPr.rFonts.set(qn("w:eastAsia"), std_heading_font)
        title_run.font.size = Pt(20)
        title_run.bold = True
        title_run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
        _set_paragraph_spacing(title_para, before=Pt(4), after=Pt(8))

    # 渲染（跳过第一个 H1 标题，避免和顶部标题重复）
    blocks = parse_markdown(content, ref_id_to_index)
    if blocks and blocks[0].type == BlockType.HEADING and blocks[0].level == 1:
        blocks = blocks[1:]
    _render_blocks(doc, blocks, font_name=std_body_font, body_size=std_body_size,
                   heading_font=std_heading_font, ref_id_to_index=ref_id_to_index,
                   heading_spacing_before=Pt(8), heading_spacing_after=Pt(4),
                   list_spacing_before=Pt(0), list_spacing_after=Pt(0),
                   block_style=NON_PREVIEW_BLOCK_STYLES["standard"])

    # 参考来源
    if references:
        _add_reference_list(doc, references, font_name=std_body_font, body_size=std_body_size)

    # 页脚
    if footer_text:
        doc.add_paragraph("")
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run(footer_text)
        footer_run.font.name = std_body_font
        footer_run._element.rPr.rFonts.set(qn("w:eastAsia"), std_body_font)
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.save(output_path)
    logger.info(f"[DocxExport] standard 格式导出完成: {output_path}")
