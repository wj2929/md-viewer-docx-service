"""
Markdown → Word 文档生成器

支持多种导出格式：
- standard: 通用格式（清晰易读，适合日常使用）
- official: 正式公文格式（GB/T 9704-2012 严格版，四级标题体系）
- internal: 机关内部文件格式
- report: 调研/分析报告格式

Markdown 解析策略：使用正则逐行解析，避免引入额外依赖。
支持：标题(H1-H4)、段落、无序/有序列表、表格、代码块、粗体/斜体。
"""
import re
import logging
from typing import List, Optional, Dict
from dataclasses import dataclass, field
from enum import Enum

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


# ==================== 数据结构 ====================

class BlockType(Enum):
    HEADING = "heading"
    PARAGRAPH = "paragraph"
    UNORDERED_LIST = "unordered_list"
    ORDERED_LIST = "ordered_list"
    TABLE = "table"
    CODE_BLOCK = "code_block"
    HORIZONTAL_RULE = "horizontal_rule"


@dataclass
class TextRun:
    """行内文本片段（支持粗体/斜体/代码/引用角标）"""
    text: str
    bold: bool = False
    italic: bool = False
    code: bool = False
    citation: bool = False       # 引用角标（上标样式）
    citation_ref_id: str = ""    # 关联的 ref_id（如 "ref_001"）


@dataclass
class Block:
    """文档块"""
    type: BlockType
    level: int = 0  # heading level (1-4)
    runs: List[TextRun] = field(default_factory=list)
    text: str = ""
    rows: List[List[str]] = field(default_factory=list)  # table rows


@dataclass
class MessageDTO:
    """消息数据传输对象"""
    role: str  # "user" | "assistant"
    content: str


# ==================== Markdown 解析 ====================

_RE_HEADING = re.compile(r"^(#{1,4})\s+(.+)$")
_RE_UL = re.compile(r"^(\s*)[-*+]\s+(.+)$")
_RE_OL = re.compile(r"^(\s*)\d+[.)]\s+(.+)$")
_RE_CODE_FENCE = re.compile(r"^```")
_RE_HR = re.compile(r"^(-{3,}|_{3,}|\*{3,})\s*$")
_RE_TABLE_ROW = re.compile(r"^\|(.+)\|$")
_RE_TABLE_SEP = re.compile(r"^\|[\s:|-]+\|$")


def parse_inline(text: str, ref_id_to_index: Optional[Dict[str, int]] = None) -> List[TextRun]:
    """解析行内格式：**bold**, *italic*, `code`, [ref_XXX] 引用角标

    Args:
        ref_id_to_index: ref_id → 显示序号 映射（如 {"ref_001": 1}）。
                         提供时 [ref_XXX] 渲染为上标 [N]；未提供时保留原文。
    """
    runs: List[TextRun] = []
    # 引用角标 + 行内格式统一正则
    pattern = re.compile(
        r"(\[(?:ref:)?(ref_\d+)(?:[^\]]*)\])"  # [ref_001] / [ref:ref_001] / [ref_006评论区]
        r"|(\[(\d{1,3})\])(?!\()"               # [6] 纯数字（排除 markdown 链接）
        r"|(\*\*\*(.+?)\*\*\*)"                 # ***bold italic***
        r"|(\*\*(.+?)\*\*)"                     # **bold**
        r"|(\*(.+?)\*)"                         # *italic*
        r"|(`(.+?)`)"                           # `code`
    )
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            runs.append(TextRun(text=text[pos:m.start()]))

        if m.group(2):  # [ref_XXX] 格式
            ref_id = m.group(2)
            if ref_id_to_index and ref_id in ref_id_to_index:
                idx = ref_id_to_index[ref_id]
                runs.append(TextRun(text=f"[{idx}]", citation=True, citation_ref_id=ref_id))
            else:
                runs.append(TextRun(text=m.group(0)))  # 无映射时保留原文
        elif m.group(4):  # [N] 纯数字
            num = int(m.group(4))
            ref_id = f"ref_{num:03d}"
            if ref_id_to_index and ref_id in ref_id_to_index:
                idx = ref_id_to_index[ref_id]
                runs.append(TextRun(text=f"[{idx}]", citation=True, citation_ref_id=ref_id))
            else:
                runs.append(TextRun(text=m.group(3)))  # 保留原文
        elif m.group(6):  # bold italic
            runs.append(TextRun(text=m.group(6), bold=True, italic=True))
        elif m.group(8):  # bold
            runs.append(TextRun(text=m.group(8), bold=True))
        elif m.group(10):  # italic
            runs.append(TextRun(text=m.group(10), italic=True))
        elif m.group(12):  # code
            runs.append(TextRun(text=m.group(12), code=True))
        pos = m.end()

    if pos < len(text):
        runs.append(TextRun(text=text[pos:]))

    return runs or [TextRun(text=text)]


def parse_markdown(md_text: str, ref_id_to_index: Optional[Dict[str, int]] = None) -> List[Block]:
    """将 Markdown 文本解析为 Block 列表"""
    blocks: List[Block] = []
    lines = md_text.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]

        # 空行跳过
        if not line.strip():
            i += 1
            continue

        # 代码块
        if _RE_CODE_FENCE.match(line.strip()):
            code_lines = []
            i += 1
            while i < len(lines) and not _RE_CODE_FENCE.match(lines[i].strip()):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing fence
            blocks.append(Block(
                type=BlockType.CODE_BLOCK,
                text="\n".join(code_lines),
            ))
            continue

        # 水平线
        if _RE_HR.match(line.strip()):
            blocks.append(Block(type=BlockType.HORIZONTAL_RULE))
            i += 1
            continue

        # 标题
        m = _RE_HEADING.match(line.strip())
        if m:
            level = len(m.group(1))
            blocks.append(Block(
                type=BlockType.HEADING,
                level=level,
                runs=parse_inline(m.group(2).strip(), ref_id_to_index),
                text=m.group(2).strip(),
            ))
            i += 1
            continue

        # 表格
        if _RE_TABLE_ROW.match(line.strip()):
            rows = []
            while i < len(lines) and _RE_TABLE_ROW.match(lines[i].strip()):
                row_text = lines[i].strip()
                if _RE_TABLE_SEP.match(row_text):
                    i += 1
                    continue
                cells = [c.strip() for c in row_text.strip("|").split("|")]
                rows.append(cells)
                i += 1
            blocks.append(Block(type=BlockType.TABLE, rows=rows))
            continue

        # 无序列表
        m = _RE_UL.match(line)
        if m:
            blocks.append(Block(
                type=BlockType.UNORDERED_LIST,
                runs=parse_inline(m.group(2), ref_id_to_index),
                text=m.group(2),
            ))
            i += 1
            continue

        # 有序列表
        m = _RE_OL.match(line)
        if m:
            blocks.append(Block(
                type=BlockType.ORDERED_LIST,
                runs=parse_inline(m.group(2), ref_id_to_index),
                text=m.group(2),
            ))
            i += 1
            continue

        # 普通段落（可能多行连续）
        para_lines = [line.strip()]
        i += 1
        while i < len(lines) and lines[i].strip() and not any([
            _RE_HEADING.match(lines[i].strip()),
            _RE_UL.match(lines[i]),
            _RE_OL.match(lines[i]),
            _RE_CODE_FENCE.match(lines[i].strip()),
            _RE_HR.match(lines[i].strip()),
            _RE_TABLE_ROW.match(lines[i].strip()),
        ]):
            para_lines.append(lines[i].strip())
            i += 1

        full_text = " ".join(para_lines)
        # 跳过 blockquote 标记 ">"
        full_text = re.sub(r"^>\s*", "", full_text)
        blocks.append(Block(
            type=BlockType.PARAGRAPH,
            runs=parse_inline(full_text, ref_id_to_index),
            text=full_text,
        ))

    return blocks


# ==================== Word 文档生成 ====================

def _apply_runs(paragraph, runs: List[TextRun], font_name: str = "宋体", font_size: Pt = Pt(12)):
    """将 TextRun 列表写入 Word 段落"""
    for run_data in runs:
        run = paragraph.add_run(run_data.text)
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        run.font.size = font_size
        if run_data.bold:
            run.bold = True
        if run_data.italic:
            run.italic = True
        if run_data.code:
            run.font.name = "Courier New"
            run.font.size = Pt(font_size.pt - 1)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        if run_data.citation:
            run.font.size = Pt(max(font_size.pt - 3, 7))
            run.font.superscript = True
            run.font.color.rgb = RGBColor(0x20, 0x80, 0xF0)


def _set_paragraph_spacing(paragraph, before: Pt = Pt(0), after: Pt = Pt(0), line: Pt = None):
    """设置段落间距"""
    pf = paragraph.paragraph_format
    pf.space_before = before
    pf.space_after = after
    if line:
        pf.line_spacing = line


def _add_table(doc, rows: List[List[str]], font_name: str = "宋体", font_size: Pt = Pt(10.5),
               ref_id_to_index: Optional[Dict[str, int]] = None):
    """添加 Word 表格（单元格内支持 **加粗**、*斜体*、`代码`、引用角标）"""
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols, style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < n_cols:
                cell = table.cell(i, j)
                cell.text = ""
                p = cell.paragraphs[0]
                # 表头行：整行加粗
                if i == 0:
                    run = p.add_run(cell_text)
                    run.font.name = font_name
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
                    run.font.size = font_size
                    run.bold = True
                else:
                    # 数据行：解析 inline markdown（加粗、斜体、引用等）
                    inline_runs = parse_inline(cell_text, ref_id_to_index)
                    _apply_runs(p, inline_runs, font_name=font_name, font_size=font_size)


def generate_standard_docx(
    title: str,
    messages: List[MessageDTO],
    output_path: str,
    footer_text: str = "由终身教育智能体生成",
):
    """
    通用格式 Word 导出

    样式：宋体正文、标题加粗、适中行距、清晰的角色分隔
    """
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # 文档标题
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    title_run.font.name = "黑体"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    title_run.font.size = Pt(18)
    title_run.bold = True
    _set_paragraph_spacing(title_para, before=Pt(12), after=Pt(12))

    # 分隔线
    doc.add_paragraph("").runs  # 空行

    # 渲染每条消息
    for msg in messages:
        # 角色标签
        role_name = "用户" if msg.role == "user" else "助手"
        role_para = doc.add_paragraph()
        role_run = role_para.add_run(f"【{role_name}】")
        role_run.font.name = "黑体"
        role_run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        role_run.font.size = Pt(12)
        role_run.bold = True
        if msg.role == "user":
            role_run.font.color.rgb = RGBColor(0x20, 0x80, 0xF0)
        else:
            role_run.font.color.rgb = RGBColor(0x18, 0xA0, 0x58)
        _set_paragraph_spacing(role_para, before=Pt(12), after=Pt(4))

        # 解析并渲染消息内容
        blocks = parse_markdown(msg.content)
        _render_blocks(doc, blocks, font_name="宋体", body_size=Pt(12))

    # 页脚
    doc.add_paragraph("")
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(footer_text)
    footer_run.font.name = "宋体"
    footer_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.save(output_path)
    logger.info(f"[DocxExport] 标准格式导出完成: {output_path}")


def generate_official_docx(
    title: str,
    messages: List[MessageDTO],
    output_path: str,
    footer_text: str = "由终身教育智能体生成",
):
    """
    行政公文标准格式 Word 导出（GB/T 9704-2012 近似）

    排版规范：
    - 标题：黑体 二号（22pt），居中
    - 正文：宋体 三号（16pt），两端对齐
    - 行距：28磅固定行距
    - 页边距：上3.7cm 下3.5cm 左2.8cm 右2.6cm
    """
    doc = Document()

    # 页面设置（公文标准页边距）
    section = doc.sections[0]
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)

    # 文档标题：黑体 二号
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    title_run.font.name = "黑体"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    title_run.font.size = Pt(22)
    title_run.bold = True
    _set_paragraph_spacing(title_para, before=Pt(0), after=Pt(16), line=Pt(28))

    # 空行分隔
    sep = doc.add_paragraph()
    _set_paragraph_spacing(sep, line=Pt(28))

    # 公文格式：只渲染 assistant 消息的内容作为正文
    # 用户消息不显示（公文是最终输出文件）
    assistant_msgs = [m for m in messages if m.role == "assistant"]
    if not assistant_msgs:
        assistant_msgs = messages  # fallback: 没有 assistant 消息时显示所有

    for msg in assistant_msgs:
        blocks = parse_markdown(msg.content)
        _render_blocks(
            doc, blocks,
            font_name="宋体",
            body_size=Pt(16),
            heading_font="黑体",
            line_spacing=Pt(28),
            first_line_indent=Cm(0.74),  # 两字符缩进
            align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        )

    # 页脚
    doc.add_paragraph("")
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(footer_text)
    footer_run.font.name = "宋体"
    footer_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.save(output_path)
    logger.info(f"[DocxExport] 公文格式导出完成: {output_path}")


def _render_blocks(
    doc: Document,
    blocks: List[Block],
    font_name: str = "宋体",
    body_size: Pt = Pt(12),
    heading_font: str = "黑体",
    line_spacing: Optional[Pt] = None,
    first_line_indent: Optional[Cm] = None,
    align: int = WD_ALIGN_PARAGRAPH.LEFT,
    heading_styles: Optional[Dict[int, "HeadingStyleDef"]] = None,
    ref_id_to_index: Optional[Dict[str, int]] = None,
):
    """将解析后的 Block 列表渲染到 Word 文档

    Args:
        heading_styles: 按级别配置标题样式 {level: HeadingStyleDef}。
                        如果提供，则覆盖 heading_font 和 heading_sizes 的线性递减逻辑。
    """

    heading_sizes = {
        1: Pt(body_size.pt + 6),
        2: Pt(body_size.pt + 4),
        3: Pt(body_size.pt + 2),
        4: Pt(body_size.pt + 1),
    }

    list_counter = 0  # 有序列表计数器
    prev_was_ol = False

    for idx, block in enumerate(blocks):
        if block.type == BlockType.HEADING:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT

            if heading_styles and block.level in heading_styles:
                hs = heading_styles[block.level]
                h_font = hs.font if isinstance(hs, HeadingStyleDef) else hs.get("font", heading_font)
                h_size = Pt(hs.size if isinstance(hs, HeadingStyleDef) else hs.get("size", body_size.pt))
                h_bold = hs.bold if isinstance(hs, HeadingStyleDef) else hs.get("bold", False)
                _apply_runs(para, block.runs, font_name=h_font, font_size=h_size)
                for run in para.runs:
                    run.bold = h_bold
            else:
                size = heading_sizes.get(block.level, body_size)
                _apply_runs(para, block.runs, font_name=heading_font, font_size=size)
                for run in para.runs:
                    run.bold = True

            _set_paragraph_spacing(para, before=Pt(12), after=Pt(6), line=line_spacing)

        elif block.type == BlockType.PARAGRAPH:
            para = doc.add_paragraph()
            para.alignment = align
            _apply_runs(para, block.runs, font_name=font_name, font_size=body_size)
            _set_paragraph_spacing(para, before=Pt(2), after=Pt(2), line=line_spacing)
            if first_line_indent:
                para.paragraph_format.first_line_indent = first_line_indent

        elif block.type == BlockType.UNORDERED_LIST:
            para = doc.add_paragraph()
            para.alignment = align
            bullet_run = para.add_run("  \u2022  ")
            bullet_run.font.name = font_name
            bullet_run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
            bullet_run.font.size = body_size
            _apply_runs(para, block.runs, font_name=font_name, font_size=body_size)
            _set_paragraph_spacing(para, before=Pt(1), after=Pt(1), line=line_spacing)
            prev_was_ol = False

        elif block.type == BlockType.ORDERED_LIST:
            if not prev_was_ol:
                list_counter = 0
            list_counter += 1
            para = doc.add_paragraph()
            para.alignment = align
            num_run = para.add_run(f"  {list_counter}.  ")
            num_run.font.name = font_name
            num_run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
            num_run.font.size = body_size
            _apply_runs(para, block.runs, font_name=font_name, font_size=body_size)
            _set_paragraph_spacing(para, before=Pt(1), after=Pt(1), line=line_spacing)
            prev_was_ol = True

        elif block.type == BlockType.TABLE:
            _add_table(doc, block.rows, font_name=font_name, font_size=Pt(body_size.pt - 1.5),
                       ref_id_to_index=ref_id_to_index)

        elif block.type == BlockType.CODE_BLOCK:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = para.add_run(block.text)
            run.font.name = "Courier New"
            run.font.size = Pt(body_size.pt - 2)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            # 代码块浅灰背景通过段落底纹实现
            shading = para._element.makeelement(qn("w:shd"), {
                qn("w:val"): "clear",
                qn("w:color"): "auto",
                qn("w:fill"): "F5F5F5",
            })
            para._element.get_or_add_pPr().append(shading)
            _set_paragraph_spacing(para, before=Pt(4), after=Pt(4), line=line_spacing)

        elif block.type == BlockType.HORIZONTAL_RULE:
            # 如果下一个 block 是标题，跳过水平线（标题自带间距，避免多余空行）
            next_block = blocks[idx + 1] if idx + 1 < len(blocks) else None
            if next_block and next_block.type == BlockType.HEADING:
                continue
            para = doc.add_paragraph()
            _set_paragraph_spacing(para, before=Pt(4), after=Pt(4))

        # 重置有序列表计数器
        if block.type != BlockType.ORDERED_LIST:
            prev_was_ol = False


# ==================== 预设体系 ====================

@dataclass
class HeadingStyleDef:
    """标题层级样式定义"""
    font: str
    size: int       # pt
    bold: bool = False


DOCX_PRESETS = {
    # ── 正式公文：GB/T 9704-2012 严格版，四级标题体系 ──
    # heading_styles 按 markdown heading level 映射（# 已被 title_font 单独处理）
    # ## = 公文一级标题（一、）→ 黑体
    # ### = 公文二级标题（（一））→ 楷体_GB2312 加粗
    # #### = 公文三级标题（1.）→ 仿宋_GB2312 加粗
    # H5 = 公文四级标题（（1））→ 仿宋_GB2312 加粗
    "official": {
        "display_name": "正式公文",
        "page_margins": {"top": 3.7, "bottom": 3.5, "left": 2.8, "right": 2.6},
        "title_font": "方正小标宋简体", "title_size": 22,
        "heading_styles": {
            1: HeadingStyleDef("方正小标宋简体", 22),          # # 文档标题（与 title_font 一致，兜底）
            2: HeadingStyleDef("黑体", 16),                   # ## 一、公文一级标题
            3: HeadingStyleDef("楷体_GB2312", 16, bold=True), # ### （一）公文二级标题
            4: HeadingStyleDef("仿宋_GB2312", 16, bold=True), # #### 1. 公文三级标题
        },
        "body_font": "仿宋_GB2312", "body_size": 16,
        "line_spacing": 28,
        "first_line_indent": 0.74,
        "align": "justify",
    },
    # ── 机关内部文件 ──
    # ## = 一级标题 → 黑体, ### = 二级标题 → 宋体加粗
    "internal": {
        "display_name": "机关内部文件",
        "page_margins": {"top": 2.54, "bottom": 2.54, "left": 3.17, "right": 3.17},
        "title_font": "黑体", "title_size": 18,
        "heading_styles": {
            1: HeadingStyleDef("黑体", 18),                    # # 文档标题（兜底）
            2: HeadingStyleDef("黑体", 15),                    # ## 一级标题
            3: HeadingStyleDef("宋体", 15, bold=True),         # ### 二级标题
            4: HeadingStyleDef("宋体", 15, bold=True),         # #### 三级标题
        },
        "body_font": "宋体", "body_size": 15,
        "line_spacing_multiple": 1.5,
        "first_line_indent": 0.74,
        "align": "justify",
    },
    # ── 调研/分析报告 ──
    # ## = 一级标题 → 黑体, ### = 二级标题 → 宋体加粗
    "report": {
        "display_name": "调研/分析报告",
        "page_margins": {"top": 2.54, "bottom": 2.54, "left": 3.17, "right": 3.17},
        "title_font": "黑体", "title_size": 16,
        "heading_styles": {
            1: HeadingStyleDef("黑体", 16),                    # # 文档标题（兜底）
            2: HeadingStyleDef("黑体", 14),                    # ## 一级标题
            3: HeadingStyleDef("宋体", 12, bold=True),         # ### 二级标题
            4: HeadingStyleDef("宋体", 12, bold=True),         # #### 三级标题
        },
        "body_font": "宋体", "body_size": 12,
        "line_spacing_multiple": 1.5,
        "first_line_indent": 0.56,
        "align": "justify",
    },
}

VALID_STYLES = frozenset(DOCX_PRESETS.keys()) | {"standard"}


def _build_ref_id_to_index(references: Optional[List[dict]]) -> Optional[Dict[str, int]]:
    """从 references 列表构建 ref_id → 显示序号 映射"""
    if not references:
        return None
    mapping: Dict[str, int] = {}
    for i, ref in enumerate(references, 1):
        ref_id = ref.get('id', '')
        if ref_id:
            mapping[ref_id] = i
    return mapping if mapping else None


def _add_reference_list(
    doc: Document,
    references: List[dict],
    font_name: str = "宋体",
    body_size: Pt = Pt(12),
    heading_font: str = "黑体",
    line_spacing: Optional[Pt] = None,
):
    """在文档末尾添加"参考来源"章节"""
    if not references:
        return

    # 分隔线
    doc.add_paragraph("")

    # 标题
    heading_para = doc.add_paragraph()
    heading_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading_run = heading_para.add_run("参考来源")
    heading_run.font.name = heading_font
    heading_run._element.rPr.rFonts.set(qn("w:eastAsia"), heading_font)
    heading_run.font.size = Pt(body_size.pt + 2)
    heading_run.bold = True
    _set_paragraph_spacing(heading_para, before=Pt(12), after=Pt(6), line=line_spacing)

    # 逐条列出
    for i, ref in enumerate(references, 1):
        title = ref.get('title', ref.get('filename', f'来源 {i}'))
        url = ref.get('url', '')
        snippet = ref.get('snippet', '')

        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # 序号
        idx_run = para.add_run(f"[{i}] ")
        idx_run.font.name = font_name
        idx_run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        idx_run.font.size = Pt(body_size.pt - 1)
        idx_run.bold = True

        # 标题
        title_run = para.add_run(title)
        title_run.font.name = font_name
        title_run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        title_run.font.size = Pt(body_size.pt - 1)

        # URL
        if url:
            url_run = para.add_run(f"  {url}")
            url_run.font.name = font_name
            url_run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
            url_run.font.size = Pt(body_size.pt - 2)
            url_run.font.color.rgb = RGBColor(0x20, 0x80, 0xF0)

        _set_paragraph_spacing(para, before=Pt(1), after=Pt(1), line=line_spacing)


def generate_docx_from_content(
    content: str,
    output_path: str,
    *,
    style: str = "standard",
    title: str = None,
    footer_text: str = "由终身教育智能体生成",
    references: Optional[List[dict]] = None,
):
    """
    从 Markdown 内容生成格式化 Word 文档（统一入口）

    Args:
        content: Markdown 格式的文档内容
        output_path: 输出文件路径
        style: 排版预设 ID（official/internal/report/standard）
        title: 文档标题，为空时自动从内容首个 # 标题提取
        footer_text: 页脚文字
        references: 引用列表（来自 ArtifactVersion.metadata_），用于渲染引用角标和参考来源
    """
    # 白名单校验
    if style not in VALID_STYLES:
        logger.warning(f"[DocxExport] 无效的 style 参数: {style!r}，回退到 standard")
        style = "standard"

    # 构建引用映射
    ref_id_to_index = _build_ref_id_to_index(references)

    # standard 模式：基础渲染
    if style == "standard":
        _generate_standard_from_content(content, output_path, title, footer_text, references=references, ref_id_to_index=ref_id_to_index)
        return

    preset = DOCX_PRESETS[style]

    # 提取标题
    if not title:
        for line in content.split('\n'):
            stripped = line.strip()
            if stripped.startswith('# '):
                title = stripped.lstrip('#').strip()
                break
        if not title:
            title = preset.get("display_name", "文档")

    doc = Document()

    # 页面设置
    section = doc.sections[0]
    margins = preset["page_margins"]
    section.top_margin = Cm(margins["top"])
    section.bottom_margin = Cm(margins["bottom"])
    section.left_margin = Cm(margins["left"])
    section.right_margin = Cm(margins["right"])

    # 文档标题
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    title_font_name = preset["title_font"]
    title_run.font.name = title_font_name
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), title_font_name)
    title_run.font.size = Pt(preset["title_size"])
    title_run.bold = True

    # 标题行距
    line_sp = Pt(preset["line_spacing"]) if "line_spacing" in preset else None
    _set_paragraph_spacing(title_para, before=Pt(0), after=Pt(16), line=line_sp)

    # 空行
    sep = doc.add_paragraph()
    _set_paragraph_spacing(sep, line=line_sp)

    # 正文渲染参数
    body_font = preset["body_font"]
    body_size = Pt(preset["body_size"])
    first_indent = Cm(preset["first_line_indent"]) if preset.get("first_line_indent") else None
    align_map = {"justify": WD_ALIGN_PARAGRAPH.JUSTIFY, "left": WD_ALIGN_PARAGRAPH.LEFT}
    align_val = align_map.get(preset.get("align", "justify"), WD_ALIGN_PARAGRAPH.JUSTIFY)

    # 行距处理：固定行距 vs 倍数行距
    if "line_spacing" in preset:
        line_spacing_val = Pt(preset["line_spacing"])
    else:
        line_spacing_val = None

    # 解析并渲染（去掉第一个 H1 标题，因为已单独渲染为文档标题）
    blocks = parse_markdown(content, ref_id_to_index)
    if blocks and blocks[0].type == BlockType.HEADING and blocks[0].level == 1:
        blocks = blocks[1:]
    _render_blocks(
        doc, blocks,
        font_name=body_font,
        body_size=body_size,
        heading_font=preset.get("title_font", "黑体"),
        line_spacing=line_spacing_val,
        first_line_indent=first_indent,
        align=align_val,
        heading_styles=preset.get("heading_styles"),
        ref_id_to_index=ref_id_to_index,
    )

    # 倍数行距需要额外处理（_render_blocks 只支持固定行距参数）
    if "line_spacing_multiple" in preset:
        from docx.shared import Emu
        for para in doc.paragraphs[2:]:  # 跳过标题和分隔段
            pf = para.paragraph_format
            if pf.line_spacing is None:
                pf.line_spacing = preset["line_spacing_multiple"]

    # 参考来源
    if references:
        _add_reference_list(
            doc, references,
            font_name=body_font, body_size=body_size,
            heading_font=preset.get("title_font", "黑体"),
            line_spacing=line_spacing_val,
        )

    # 页脚
    doc.add_paragraph("")
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(footer_text)
    footer_run.font.name = "宋体"
    footer_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.save(output_path)
    logger.info(f"[DocxExport] {style} 格式导出完成: {output_path}")


def _generate_standard_from_content(
    content: str,
    output_path: str,
    title: str = None,
    footer_text: str = "由终身教育智能体生成",
    references: Optional[List[dict]] = None,
    ref_id_to_index: Optional[Dict[str, int]] = None,
):
    """标准格式渲染（基础 markdown→word，与原 FileService._write_docx 等效）"""
    if not title:
        for line in content.split('\n'):
            stripped = line.strip()
            if stripped.startswith('# '):
                title = stripped.lstrip('#').strip()
                break

    doc = Document()

    # 标准页面设置
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # 标题（如果有）
    if title:
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(title)
        title_run.font.name = "黑体"
        title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        title_run.font.size = Pt(18)
        title_run.bold = True
        _set_paragraph_spacing(title_para, before=Pt(12), after=Pt(12))
        doc.add_paragraph("")

    # 渲染（跳过第一个 H1 标题，避免和顶部标题重复）
    blocks = parse_markdown(content, ref_id_to_index)
    if blocks and blocks[0].type == BlockType.HEADING and blocks[0].level == 1:
        blocks = blocks[1:]
    _render_blocks(doc, blocks, font_name="宋体", body_size=Pt(12), ref_id_to_index=ref_id_to_index)

    # 参考来源
    if references:
        _add_reference_list(doc, references, font_name="宋体", body_size=Pt(12))

    # 页脚
    if footer_text:
        doc.add_paragraph("")
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run(footer_text)
        footer_run.font.name = "宋体"
        footer_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.save(output_path)
    logger.info(f"[DocxExport] standard 格式导出完成: {output_path}")
